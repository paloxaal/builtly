  from __future__ import annotations

import io
import json
import math
import sys
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence
import xml.etree.ElementTree as ET

import pandas as pd
import streamlit as st

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

try:
    from builtly_ai_fallback import generate_json_with_fallback
except Exception:
    generate_json_with_fallback = None

from builtly_module_kit import (
    configure_page,
    dataframe_download,
    json_download,
    render_hero,
    render_metric_cards,
    render_panel,
    render_project_snapshot,
    render_section,
    tone_from_score,
)

try:
    from builtly_public_data import geocode_address
except Exception:
    def geocode_address(address: str, municipality: str = "") -> Dict[str, Any]:
        return {"status": "missing", "source": "Address resolver", "note": "No live geocoder configured"}

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    REPORTLAB_READY = True
except Exception:
    REPORTLAB_READY = False

# ---------------------------------------------------------
# Localisation
# ---------------------------------------------------------

LANG = {
    "no": {
        "page_title": "Builtly | Klimarisiko",
        "eyebrow": "Klimarisiko",
        "title": "Screen klimarisiko for eiendommer og porteføljer.",
        "subtitle": "Vurder flom, skred, havnivå og varmestress i én arbeidsflyt, og generer rapportpakker som kan deles videre med investorer, banker og prosjektteam.",
        "setup_kicker": "Oppsett",
        "setup_title": "Hva skal analyseres?",
        "setup_sub": "Velg språk og marked først. Deretter kan du analysere én eiendom eller laste opp en portefølje.",
        "lang_label": "Språk",
        "market_label": "Marked / jurisdiksjon",
        "mode_label": "Analysemodus",
        "mode_single": "Enkeltobjekt",
        "mode_portfolio": "Portefølje",
        "property_name": "Eiendomsnavn",
        "address": "Adresse / lokasjon",
        "municipality": "Kommune / by",
        "region": "Region / delstat / county",
        "asset_class": "Eiendomstype",
        "year_built": "Byggeår",
        "area": "Areal (m²)",
        "value": "Anslått eiendomsverdi",
        "currency_hint": "Valuta",
        "analysis_settings": "Scenarier og risikoparametre",
        "scenario": "Klimascenario",
        "horizon": "Tidshorisont",
        "elevation": "Høyde over havet (m)",
        "coast": "Avstand til kyst (km)",
        "river": "Avstand til elv/bekk (km)",
        "slope": "Terrenghelning (grader)",
        "soil": "Grunnforhold",
        "heat": "Urban varmestress (0–10)",
        "flood_zone": "Ligger i eller nær flomsone",
        "landslide_zone": "Ligger i eller nær skred-/rasområde",
        "basement": "Bygget har kjeller / underetasje",
        "upload_portfolio": "Last opp portefølje (CSV/XLSX)",
        "weights_title": "Vektsetting for risikofaktorer",
        "weights_help": "Vektene styrer den samlede klimarisikoscoren. Standardvektene følger produktspesifikasjonen.",
        "generate": "Generer klimarisikorapport",
        "rerun": "Kjør analysen på nytt",
        "need_inputs": "Legg inn adresse eller koordinater, eller last opp en portefølje, før du genererer rapport.",
        "results_kicker": "Resultater",
        "results_title": "Resultater og sporbarhet",
        "results_sub": "Du får både menneskelesbar rapport, maskinlesbar eksport og regulatorisk mapping.",
        "metrics_score": "Klimarisikoscore",
        "metrics_uncertainty": "Usikkerhet",
        "metrics_assets": "Eiendommer",
        "metrics_framework": "Primært rapporteringsspor",
        "tab_overview": "Oversikt",
        "tab_factors": "Risikofaktorer",
        "tab_mapping": "Regulatorisk mapping",
        "tab_sources": "Datakilder",
        "tab_package": "Rapportpakke",
        "tab_portfolio": "Portefølje",
        "summary_title": "Kort oppsummering",
        "drivers_title": "Hva som driver risikoen",
        "actions_title": "Anbefalte neste steg",
        "factors_title": "Faktorscore",
        "sources_title": "Kilder og datadekning",
        "mapping_title": "Regulatorisk mapping",
        "report_title": "Rapportpakke",
        "report_sub": "Når rapporten er generert kan du laste ned PDF, DOCX, JSON, XML og CSV.",
        "disclaimer_ack": "Jeg forstår at dette er et AI-generert dataprodukt og ikke en faglig attestasjon.",
        "downloads_locked": "Bekreft forståelsen av disclaimere før du laster ned rapportpakken.",
        "download_pdf": "Last ned PDF-rapport",
        "download_docx": "Last ned DOCX-rapport",
        "download_json": "Last ned JSON",
        "download_xml": "Last ned XML-mapping",
        "download_csv": "Last ned CSV med faktorscore",
        "download_portfolio": "Last ned portefølje-CSV",
        "portfolio_summary": "Porteføljesammendrag",
        "portfolio_placeholder": "Bytt til portefølje for å analysere mange eiendommer samtidig.",
        "live_data_title": "Markedstilpasning",
        "live_data_sub": "Systemet bruker valgt land til å vise riktig regelverk, disclaimer og datakilder. Der live API ikke er koblet, brukes strukturerte input og predefinert datakatalog.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Rapportpakke generert.",
        "ai_unavailable": "AI-sammendrag er ikke tilgjengelig akkurat nå. Builtly viser en strukturert standardoppsummering i stedet.",
        "map_hint": "Kartbilde kan legges til når markedets GIS-kilde er koblet live eller når koordinater og faresoner er tilgjengelige i tenant-oppsettet.",
        "top_risk": "Høyest risikofaktor",
        "risk_band": "Risikoklasse",
        "country_note_label": "Juridisk/jurisdiksjonell merknad",
    },
    "sv": {
        "page_title": "Builtly | Klimatrisk",
        "eyebrow": "Klimatrisk",
        "title": "Screena klimatrisk för enstaka fastigheter och hela portföljer.",
        "subtitle": "Bedöm översvämning, skred, havsnivå och värmestress i ett arbetsflöde och generera rapportpaket som kan delas med investerare, banker och projektteam.",
        "setup_kicker": "Inställning",
        "setup_title": "Vad ska analyseras?",
        "setup_sub": "Välj språk och marknad först. Därefter kan du analysera en fastighet eller ladda upp en portfölj.",
        "lang_label": "Språk",
        "market_label": "Marknad / jurisdiktion",
        "mode_label": "Analysläge",
        "mode_single": "Enskild fastighet",
        "mode_portfolio": "Portfölj",
        "property_name": "Fastighetsnamn",
        "address": "Adress / plats",
        "municipality": "Kommun / stad",
        "region": "Region / delstat / county",
        "asset_class": "Fastighetstyp",
        "year_built": "Byggår",
        "area": "Area (m²)",
        "value": "Uppskattat fastighetsvärde",
        "currency_hint": "Valuta",
        "analysis_settings": "Scenarier och riskparametrar",
        "scenario": "Klimatscenario",
        "horizon": "Tidshorisont",
        "elevation": "Höjd över havet (m)",
        "coast": "Avstånd till kust (km)",
        "river": "Avstånd till älv/bäck (km)",
        "slope": "Marklutning (grader)",
        "soil": "Markförhållanden",
        "heat": "Urban värmestress (0–10)",
        "flood_zone": "Ligger i eller nära översvämningszon",
        "landslide_zone": "Ligger i eller nära ras-/skredområde",
        "basement": "Byggnaden har källare / underplan",
        "upload_portfolio": "Ladda upp portfölj (CSV/XLSX)",
        "weights_title": "Viktning av riskfaktorer",
        "weights_help": "Vikterna styr den sammanlagda klimatriskpoängen. Standardvikterna följer produktspecifikationen.",
        "generate": "Generera klimatriskrapport",
        "rerun": "Kör analysen igen",
        "need_inputs": "Fyll i adress eller koordinater, eller ladda upp en portfölj, innan rapporten genereras.",
        "results_kicker": "Resultat",
        "results_title": "Resultat och spårbarhet",
        "results_sub": "Du får både människoläsbar rapport, maskinläsbar export och regulatorisk mapping.",
        "metrics_score": "Klimatriskpoäng",
        "metrics_uncertainty": "Osäkerhet",
        "metrics_assets": "Fastigheter",
        "metrics_framework": "Primär rapporteringsram",
        "tab_overview": "Översikt",
        "tab_factors": "Riskfaktorer",
        "tab_mapping": "Regulatorisk mapping",
        "tab_sources": "Datakällor",
        "tab_package": "Rapportpaket",
        "tab_portfolio": "Portfölj",
        "summary_title": "Kort sammanfattning",
        "drivers_title": "Vad driver risken",
        "actions_title": "Rekommenderade nästa steg",
        "factors_title": "Faktorscore",
        "sources_title": "Källor och datatäckning",
        "mapping_title": "Regulatorisk mapping",
        "report_title": "Rapportpaket",
        "report_sub": "När rapporten är genererad kan du ladda ned PDF, DOCX, JSON, XML och CSV.",
        "disclaimer_ack": "Jag förstår att detta är en AI-genererad dataprodukt och inte en fackmässig attest.",
        "downloads_locked": "Bekräfta disclaimern innan du laddar ned rapportpaketet.",
        "download_pdf": "Ladda ned PDF-rapport",
        "download_docx": "Ladda ned DOCX-rapport",
        "download_json": "Ladda ned JSON",
        "download_xml": "Ladda ned XML-mapping",
        "download_csv": "Ladda ned CSV med faktorscore",
        "download_portfolio": "Ladda ned portfölj-CSV",
        "portfolio_summary": "Portföljsammanfattning",
        "portfolio_placeholder": "Byt till portfölj för att analysera många fastigheter samtidigt.",
        "live_data_title": "Marknadsanpassning",
        "live_data_sub": "Systemet använder valt land för rätt regelverk, disclaimer och datakällor. Där live-API saknas används strukturerade indata och definierad datakatalog.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Rapportpaket genererat.",
        "ai_unavailable": "AI-sammanfattning är inte tillgänglig just nu. Builtly visar en strukturerad standardsammanfattning i stället.",
        "map_hint": "Kartbild kan läggas till när marknadens GIS-källa är kopplad live eller när koordinater och farozoner finns i tenant-konfigurationen.",
        "top_risk": "Högsta riskfaktor",
        "risk_band": "Riskklass",
        "country_note_label": "Juridisk/jurisdiktionell kommentar",
    },
    "da": {
        "page_title": "Builtly | Klimarisiko",
        "eyebrow": "Klimarisiko",
        "title": "Screen klimarisiko for enkelte ejendomme og hele porteføljer.",
        "subtitle": "Vurder oversvømmelse, skred, havniveau og varmestress i ét workflow og generér rapportpakker, som kan deles med investorer, banker og projektteams.",
        "setup_kicker": "Opsætning",
        "setup_title": "Hvad skal analyseres?",
        "setup_sub": "Vælg først sprog og marked. Derefter kan du analysere én ejendom eller uploade en portefølje.",
        "lang_label": "Sprog",
        "market_label": "Marked / jurisdiktion",
        "mode_label": "Analysemodus",
        "mode_single": "Enkelt ejendom",
        "mode_portfolio": "Portefølje",
        "property_name": "Ejendomsnavn",
        "address": "Adresse / lokation",
        "municipality": "Kommune / by",
        "region": "Region / delstat / county",
        "asset_class": "Ejendomstype",
        "year_built": "Byggeår",
        "area": "Areal (m²)",
        "value": "Anslået ejendomsværdi",
        "currency_hint": "Valuta",
        "analysis_settings": "Scenarier og risikoparametre",
        "scenario": "Klimascenarie",
        "horizon": "Tidshorisont",
        "elevation": "Højde over havet (m)",
        "coast": "Afstand til kyst (km)",
        "river": "Afstand til å / vandløb (km)",
        "slope": "Terrænhældning (grader)",
        "soil": "Jordbundsforhold",
        "heat": "Urban varmestress (0–10)",
        "flood_zone": "Ligger i eller nær oversvømmelseszone",
        "landslide_zone": "Ligger i eller nær skred-/rasområde",
        "basement": "Bygningen har kælder / underetage",
        "upload_portfolio": "Upload portefølje (CSV/XLSX)",
        "weights_title": "Vægtning af risikofaktorer",
        "weights_help": "Vægtene styrer den samlede klimarisikoscore. Standardvægtene følger produktspecifikationen.",
        "generate": "Generér klimarisikorapport",
        "rerun": "Kør analysen igen",
        "need_inputs": "Indtast adresse eller koordinater, eller upload en portefølje, før rapporten genereres.",
        "results_kicker": "Resultater",
        "results_title": "Resultater og sporbarhed",
        "results_sub": "Du får både menneskelæselig rapport, maskinlæsbar eksport og regulatorisk mapping.",
        "metrics_score": "Klimarisikoscore",
        "metrics_uncertainty": "Usikkerhed",
        "metrics_assets": "Ejendomme",
        "metrics_framework": "Primær rapporteringsramme",
        "tab_overview": "Overblik",
        "tab_factors": "Risikofaktorer",
        "tab_mapping": "Regulatorisk mapping",
        "tab_sources": "Datakilder",
        "tab_package": "Rapportpakke",
        "tab_portfolio": "Portefølje",
        "summary_title": "Kort opsummering",
        "drivers_title": "Hvad driver risikoen",
        "actions_title": "Anbefalede næste skridt",
        "factors_title": "Faktorscore",
        "sources_title": "Kilder og datadækning",
        "mapping_title": "Regulatorisk mapping",
        "report_title": "Rapportpakke",
        "report_sub": "Når rapporten er genereret kan du downloade PDF, DOCX, JSON, XML og CSV.",
        "disclaimer_ack": "Jeg forstår, at dette er et AI-genereret dataprodukt og ikke en faglig attestering.",
        "downloads_locked": "Bekræft disclaimern før du downloader rapportpakken.",
        "download_pdf": "Download PDF-rapport",
        "download_docx": "Download DOCX-rapport",
        "download_json": "Download JSON",
        "download_xml": "Download XML-mapping",
        "download_csv": "Download CSV med faktorscore",
        "download_portfolio": "Download portefølje-CSV",
        "portfolio_summary": "Porteføljesammendrag",
        "portfolio_placeholder": "Skift til portefølje for at analysere mange ejendomme samtidigt.",
        "live_data_title": "Markedstilpasning",
        "live_data_sub": "Systemet bruger valgt land til korrekt regelsæt, disclaimer og datakilder. Hvor live-API ikke er koblet, bruges strukturerede input og en defineret datakatalog.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Rapportpakke genereret.",
        "ai_unavailable": "AI-opsummering er ikke tilgængelig lige nu. Builtly viser i stedet en struktureret standardopsummering.",
        "map_hint": "Kortbillede kan tilføjes, når markedets GIS-kilde er koblet live, eller når koordinater og farezoner findes i tenant-konfigurationen.",
        "top_risk": "Højeste risikofaktor",
        "risk_band": "Risikoklasse",
        "country_note_label": "Juridisk/jurisdiktionel note",
    },
    "en-GB": {
        "page_title": "Builtly | Climate Risk",
        "eyebrow": "Climate Risk",
        "title": "Screen climate risk for single assets and full portfolios.",
        "subtitle": "Assess flood, landslide, sea-level and heat stress in one workflow and generate report packs you can share with investors, banks and project teams.",
        "setup_kicker": "Setup",
        "setup_title": "What should be analysed?",
        "setup_sub": "Select language and market first. Then analyse one asset or upload a full portfolio.",
        "lang_label": "Language",
        "market_label": "Market / jurisdiction",
        "mode_label": "Analysis mode",
        "mode_single": "Single asset",
        "mode_portfolio": "Portfolio",
        "property_name": "Asset name",
        "address": "Address / location",
        "municipality": "City / municipality",
        "region": "Region / state / county",
        "asset_class": "Asset class",
        "year_built": "Year built",
        "area": "Area (m²)",
        "value": "Estimated asset value",
        "currency_hint": "Currency",
        "analysis_settings": "Scenarios and risk parameters",
        "scenario": "Climate scenario",
        "horizon": "Time horizon",
        "elevation": "Elevation above sea level (m)",
        "coast": "Distance to coast (km)",
        "river": "Distance to river/stream (km)",
        "slope": "Slope (degrees)",
        "soil": "Ground conditions",
        "heat": "Urban heat stress (0–10)",
        "flood_zone": "Located in or near flood zone",
        "landslide_zone": "Located in or near landslide zone",
        "basement": "Building has basement / lower ground floor",
        "upload_portfolio": "Upload portfolio (CSV/XLSX)",
        "weights_title": "Risk-factor weighting",
        "weights_help": "The weights drive the total climate-risk score. Default weights follow the product specification.",
        "generate": "Generate climate risk report",
        "rerun": "Run analysis again",
        "need_inputs": "Enter an address or coordinates, or upload a portfolio, before generating the report.",
        "results_kicker": "Results",
        "results_title": "Results and traceability",
        "results_sub": "You get a human-readable report, machine-readable exports and jurisdiction-aware regulatory mapping.",
        "metrics_score": "Climate risk score",
        "metrics_uncertainty": "Uncertainty",
        "metrics_assets": "Assets",
        "metrics_framework": "Primary reporting track",
        "tab_overview": "Overview",
        "tab_factors": "Risk factors",
        "tab_mapping": "Regulatory mapping",
        "tab_sources": "Data sources",
        "tab_package": "Report pack",
        "tab_portfolio": "Portfolio",
        "summary_title": "Short summary",
        "drivers_title": "Main risk drivers",
        "actions_title": "Recommended next steps",
        "factors_title": "Factor scores",
        "sources_title": "Sources and data coverage",
        "mapping_title": "Regulatory mapping",
        "report_title": "Report pack",
        "report_sub": "Once the report has been generated you can download PDF, DOCX, JSON, XML and CSV.",
        "disclaimer_ack": "I understand that this is an AI-generated data product and not a professional attestation.",
        "downloads_locked": "Confirm the disclaimer before downloading the report pack.",
        "download_pdf": "Download PDF report",
        "download_docx": "Download DOCX report",
        "download_json": "Download JSON",
        "download_xml": "Download XML mapping",
        "download_csv": "Download CSV with factor scores",
        "download_portfolio": "Download portfolio CSV",
        "portfolio_summary": "Portfolio summary",
        "portfolio_placeholder": "Switch to portfolio mode to analyse many assets at once.",
        "live_data_title": "Market adaptation",
        "live_data_sub": "The system uses your selected market to apply the right rule set, disclaimer and source catalogue. Where live APIs are not connected yet, Builtly falls back to structured inputs and the configured source register.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Report pack generated.",
        "ai_unavailable": "AI summary is not available right now. Builtly is showing a structured default summary instead.",
        "map_hint": "A static map can be added when the market-specific GIS source is connected live or when coordinates and hazard layers are available in tenant configuration.",
        "top_risk": "Highest risk factor",
        "risk_band": "Risk band",
        "country_note_label": "Jurisdiction note",
    },
    "en-US": {
        "page_title": "Builtly | Climate Risk",
        "eyebrow": "Climate Risk",
        "title": "Screen climate risk for individual assets and full portfolios.",
        "subtitle": "Assess flood, landslide, sea-level, and heat stress in one workflow and generate report packages you can share with lenders, investors, and project teams.",
        "setup_kicker": "Setup",
        "setup_title": "What should be analyzed?",
        "setup_sub": "Select language and market first. Then analyze one asset or upload a full portfolio.",
        "lang_label": "Language",
        "market_label": "Market / jurisdiction",
        "mode_label": "Analysis mode",
        "mode_single": "Single asset",
        "mode_portfolio": "Portfolio",
        "property_name": "Asset name",
        "address": "Address / location",
        "municipality": "City / municipality",
        "region": "Region / state / county",
        "asset_class": "Asset class",
        "year_built": "Year built",
        "area": "Area (m²)",
        "value": "Estimated asset value",
        "currency_hint": "Currency",
        "analysis_settings": "Scenarios and risk parameters",
        "scenario": "Climate scenario",
        "horizon": "Time horizon",
        "elevation": "Elevation above sea level (m)",
        "coast": "Distance to coast (km)",
        "river": "Distance to river/stream (km)",
        "slope": "Slope (degrees)",
        "soil": "Ground conditions",
        "heat": "Urban heat stress (0–10)",
        "flood_zone": "Located in or near flood zone",
        "landslide_zone": "Located in or near landslide zone",
        "basement": "Building has basement / lower level",
        "upload_portfolio": "Upload portfolio (CSV/XLSX)",
        "weights_title": "Risk-factor weighting",
        "weights_help": "The weights drive the total climate-risk score. Default weights follow the product specification.",
        "generate": "Generate climate risk report",
        "rerun": "Run analysis again",
        "need_inputs": "Enter an address or coordinates, or upload a portfolio, before generating the report.",
        "results_kicker": "Results",
        "results_title": "Results and traceability",
        "results_sub": "You get a human-readable report, machine-readable exports, and jurisdiction-aware regulatory mapping.",
        "metrics_score": "Climate risk score",
        "metrics_uncertainty": "Uncertainty",
        "metrics_assets": "Assets",
        "metrics_framework": "Primary reporting track",
        "tab_overview": "Overview",
        "tab_factors": "Risk factors",
        "tab_mapping": "Regulatory mapping",
        "tab_sources": "Data sources",
        "tab_package": "Report package",
        "tab_portfolio": "Portfolio",
        "summary_title": "Short summary",
        "drivers_title": "Main risk drivers",
        "actions_title": "Recommended next steps",
        "factors_title": "Factor scores",
        "sources_title": "Sources and data coverage",
        "mapping_title": "Regulatory mapping",
        "report_title": "Report package",
        "report_sub": "Once the report has been generated you can download PDF, DOCX, JSON, XML, and CSV.",
        "disclaimer_ack": "I understand that this is an AI-generated data product and not a professional attestation.",
        "downloads_locked": "Confirm the disclaimer before downloading the report package.",
        "download_pdf": "Download PDF report",
        "download_docx": "Download DOCX report",
        "download_json": "Download JSON",
        "download_xml": "Download XML mapping",
        "download_csv": "Download CSV with factor scores",
        "download_portfolio": "Download portfolio CSV",
        "portfolio_summary": "Portfolio summary",
        "portfolio_placeholder": "Switch to portfolio mode to analyze many assets at once.",
        "live_data_title": "Market adaptation",
        "live_data_sub": "The system uses your selected market to apply the right rule set, disclaimer, and source catalog. Where live APIs are not connected yet, Builtly falls back to structured inputs and the configured source register.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Report package generated.",
        "ai_unavailable": "AI summary is not available right now. Builtly is showing a structured default summary instead.",
        "map_hint": "A static map can be added when the market-specific GIS source is connected live or when coordinates and hazard layers are available in tenant configuration.",
        "top_risk": "Highest risk factor",
        "risk_band": "Risk band",
        "country_note_label": "Jurisdiction note",
    },
    "de": {
        "page_title": "Builtly | Klimarisiko",
        "eyebrow": "Klimarisiko",
        "title": "Prüfen Sie Klimarisiken für Einzelobjekte und ganze Portfolios.",
        "subtitle": "Bewerten Sie Hochwasser, Rutschung, Meeresspiegel und Hitzestress in einem Workflow und erzeugen Sie Berichtspakete für Investoren, Banken und Projektteams.",
        "setup_kicker": "Setup",
        "setup_title": "Was soll analysiert werden?",
        "setup_sub": "Wählen Sie zuerst Sprache und Markt. Danach können Sie ein Objekt analysieren oder ein Portfolio hochladen.",
        "lang_label": "Sprache",
        "market_label": "Markt / Jurisdiktion",
        "mode_label": "Analysemodus",
        "mode_single": "Einzelobjekt",
        "mode_portfolio": "Portfolio",
        "property_name": "Objektname",
        "address": "Adresse / Standort",
        "municipality": "Gemeinde / Stadt",
        "region": "Region / Bundesland / County",
        "asset_class": "Objektart",
        "year_built": "Baujahr",
        "area": "Fläche (m²)",
        "value": "Geschätzter Immobilienwert",
        "currency_hint": "Währung",
        "analysis_settings": "Szenarien und Risikoparameter",
        "scenario": "Klimaszenario",
        "horizon": "Zeithorizont",
        "elevation": "Höhe über Meeresspiegel (m)",
        "coast": "Entfernung zur Küste (km)",
        "river": "Entfernung zu Fluss/Bach (km)",
        "slope": "Hangneigung (Grad)",
        "soil": "Bodenverhältnisse",
        "heat": "Urbaner Hitzestress (0–10)",
        "flood_zone": "Liegt in oder nahe einer Hochwasserzone",
        "landslide_zone": "Liegt in oder nahe einer Rutschungs-/Gefahrenzone",
        "basement": "Gebäude hat Keller / Untergeschoss",
        "upload_portfolio": "Portfolio hochladen (CSV/XLSX)",
        "weights_title": "Gewichtung der Risikofaktoren",
        "weights_help": "Die Gewichte steuern den Gesamtscore. Die Standardgewichte folgen der Produktspezifikation.",
        "generate": "Klimarisikobericht generieren",
        "rerun": "Analyse erneut starten",
        "need_inputs": "Geben Sie Adresse oder Koordinaten ein oder laden Sie ein Portfolio hoch, bevor der Bericht erzeugt wird.",
        "results_kicker": "Ergebnisse",
        "results_title": "Ergebnisse und Nachvollziehbarkeit",
        "results_sub": "Sie erhalten einen lesbaren Bericht, maschinenlesbare Exporte und ein marktbezogenes regulatorisches Mapping.",
        "metrics_score": "Klimarisikoscore",
        "metrics_uncertainty": "Unsicherheit",
        "metrics_assets": "Objekte",
        "metrics_framework": "Primärer Reporting-Pfad",
        "tab_overview": "Überblick",
        "tab_factors": "Risikofaktoren",
        "tab_mapping": "Regulatorisches Mapping",
        "tab_sources": "Datenquellen",
        "tab_package": "Berichtspaket",
        "tab_portfolio": "Portfolio",
        "summary_title": "Kurze Zusammenfassung",
        "drivers_title": "Haupttreiber des Risikos",
        "actions_title": "Empfohlene nächste Schritte",
        "factors_title": "Faktorscores",
        "sources_title": "Quellen und Datenabdeckung",
        "mapping_title": "Regulatorisches Mapping",
        "report_title": "Berichtspaket",
        "report_sub": "Nach der Generierung können Sie PDF, DOCX, JSON, XML und CSV herunterladen.",
        "disclaimer_ack": "Ich verstehe, dass dies ein KI-generiertes Datenprodukt und keine fachliche Attestierung ist.",
        "downloads_locked": "Bestätigen Sie den Disclaimer, bevor Sie das Berichtspaket herunterladen.",
        "download_pdf": "PDF-Bericht herunterladen",
        "download_docx": "DOCX-Bericht herunterladen",
        "download_json": "JSON herunterladen",
        "download_xml": "XML-Mapping herunterladen",
        "download_csv": "CSV mit Faktorscores herunterladen",
        "download_portfolio": "Portfolio-CSV herunterladen",
        "portfolio_summary": "Portfolio-Zusammenfassung",
        "portfolio_placeholder": "Wechseln Sie in den Portfoliomodus, um viele Objekte gleichzeitig zu analysieren.",
        "live_data_title": "Marktanpassung",
        "live_data_sub": "Das System nutzt den gewählten Markt für das richtige Regelwerk, den passenden Disclaimer und die Quellenkataloge. Wo Live-APIs noch nicht verbunden sind, arbeitet Builtly mit strukturierten Eingaben und dem konfigurierten Quellenregister.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit Trail",
        "analysis_ready": "Berichtspaket generiert.",
        "ai_unavailable": "Die KI-Zusammenfassung ist derzeit nicht verfügbar. Builtly zeigt stattdessen eine strukturierte Standardzusammenfassung an.",
        "map_hint": "Eine statische Karte kann ergänzt werden, sobald die marktbezogene GIS-Quelle live verbunden ist oder Koordinaten und Gefahrenlayer in der Tenant-Konfiguration vorliegen.",
        "top_risk": "Höchster Risikofaktor",
        "risk_band": "Risikoklasse",
        "country_note_label": "Jurisdiktionshinweis",
    },
    "fi": {
        "page_title": "Builtly | Ilmastoriski",
        "eyebrow": "Ilmastoriski",
        "title": "Arvioi ilmastoriski yksittäisille kohteille ja kokonaisille portfolioille.",
        "subtitle": "Arvioi tulva-, maanvyöry-, merenpinta- ja kuumuusriskit yhdessä työnkulussa ja luo raporttipaketteja sijoittajille, pankeille ja projektitiimeille.",
        "setup_kicker": "Asetukset",
        "setup_title": "Mitä analysoidaan?",
        "setup_sub": "Valitse ensin kieli ja markkina. Sen jälkeen voit analysoida yhden kohteen tai ladata portfolion.",
        "lang_label": "Kieli",
        "market_label": "Markkina / toimivalta",
        "mode_label": "Analyysitila",
        "mode_single": "Yksittäinen kohde",
        "mode_portfolio": "Portfolio",
        "property_name": "Kohteen nimi",
        "address": "Osoite / sijainti",
        "municipality": "Kunta / kaupunki",
        "region": "Alue / osavaltio / county",
        "asset_class": "Kohdetyyppi",
        "year_built": "Rakennusvuosi",
        "area": "Pinta-ala (m²)",
        "value": "Arvioitu kohteen arvo",
        "currency_hint": "Valuutta",
        "analysis_settings": "Skenaariot ja riskiparametrit",
        "scenario": "Ilmastoskenaario",
        "horizon": "Aikahorisontti",
        "elevation": "Korkeus merenpinnasta (m)",
        "coast": "Etäisyys rannikosta (km)",
        "river": "Etäisyys jokeen/puroon (km)",
        "slope": "Maaston kaltevuus (astetta)",
        "soil": "Maaperäolosuhteet",
        "heat": "Kaupunkilämpöstressi (0–10)",
        "flood_zone": "Sijaitsee tulvavyöhykkeellä tai sen lähellä",
        "landslide_zone": "Sijaitsee sortuma-/maanvyöryriskialueella tai sen lähellä",
        "basement": "Rakennuksessa on kellari / alakerros",
        "upload_portfolio": "Lataa portfolio (CSV/XLSX)",
        "weights_title": "Riskitekijöiden painotus",
        "weights_help": "Painot ohjaavat kokonaisilmastoriskipistettä. Oletuspainot seuraavat tuotespesifikaatiota.",
        "generate": "Luo ilmastoriskiraportti",
        "rerun": "Suorita analyysi uudelleen",
        "need_inputs": "Anna osoite tai koordinaatit tai lataa portfolio ennen raportin luontia.",
        "results_kicker": "Tulokset",
        "results_title": "Tulokset ja jäljitettävyys",
        "results_sub": "Saat ihmisen luettavan raportin, koneellisesti luettavat viennit ja markkinakohtaisen sääntelykartoituksen.",
        "metrics_score": "Ilmastoriskipiste",
        "metrics_uncertainty": "Epävarmuus",
        "metrics_assets": "Kohteet",
        "metrics_framework": "Ensisijainen raportointikehys",
        "tab_overview": "Yleiskuva",
        "tab_factors": "Riskitekijät",
        "tab_mapping": "Sääntelykartoitus",
        "tab_sources": "Tietolähteet",
        "tab_package": "Raporttipaketti",
        "tab_portfolio": "Portfolio",
        "summary_title": "Lyhyt yhteenveto",
        "drivers_title": "Riskin pääajurit",
        "actions_title": "Suositellut seuraavat vaiheet",
        "factors_title": "Tekijäpisteet",
        "sources_title": "Lähteet ja kattavuus",
        "mapping_title": "Sääntelykartoitus",
        "report_title": "Raporttipaketti",
        "report_sub": "Kun raportti on luotu, voit ladata PDF-, DOCX-, JSON-, XML- ja CSV-tiedostot.",
        "disclaimer_ack": "Ymmärrän, että tämä on tekoälyn tuottama datapohjainen tuote eikä ammatillinen todistus.",
        "downloads_locked": "Vahvista vastuuvapauslauseke ennen raporttipaketin lataamista.",
        "download_pdf": "Lataa PDF-raportti",
        "download_docx": "Lataa DOCX-raportti",
        "download_json": "Lataa JSON",
        "download_xml": "Lataa XML-kartoitus",
        "download_csv": "Lataa CSV riskitekijäpisteillä",
        "download_portfolio": "Lataa portfolio-CSV",
        "portfolio_summary": "Portfolioyhteenveto",
        "portfolio_placeholder": "Vaihda portfoliotilaan analysoidaksesi useita kohteita kerralla.",
        "live_data_title": "Markkinakohtainen sovitus",
        "live_data_sub": "Järjestelmä käyttää valittua markkinaa oikeaan sääntöpohjaan, disclaimeriin ja tietolähdeluetteloon. Jos live-API ei ole vielä käytössä, Builtly käyttää rakenteistettua syötettä ja määriteltyä lähderekisteriä.",
        "disclaimer_title": "Disclaimer",
        "audit_title": "Audit trail",
        "analysis_ready": "Raporttipaketti luotu.",
        "ai_unavailable": "Tekoäly-yhteenveto ei ole juuri nyt saatavilla. Builtly näyttää sen sijaan rakenteisen oletusyhteenvedon.",
        "map_hint": "Staattinen kartta voidaan lisätä, kun markkinakohtainen GIS-lähde on kytketty tai kun koordinaatit ja vaaratasot ovat tenant-konfiguraatiossa.",
        "top_risk": "Korkein riskitekijä",
        "risk_band": "Riskiluokka",
        "country_note_label": "Jurisdiktiohuomio",
    },
}

SOIL_OPTIONS = {
    "no": ["Morene", "Leire", "Berg", "Marine avsetninger", "Fyllmasser"],
    "sv": ["Morän", "Lera", "Berg", "Marina sediment", "Fyllnadsmassor"],
    "da": ["Moræne", "Ler", "Berg", "Marine aflejringer", "Fyldmasser"],
    "en-GB": ["Moraine", "Clay", "Bedrock", "Marine sediments", "Made ground"],
    "en-US": ["Moraine", "Clay", "Bedrock", "Marine sediments", "Made ground"],
    "de": ["Moräne", "Lehm", "Fels", "Marine Sedimente", "Auffüllung"],
    "fi": ["Moreeni", "Savi", "Kallio", "Meri­sedimentti", "Täyttömaa"],
}

ASSET_CLASSES = {
    "no": ["Bolig", "Kontor", "Logistikk", "Hotell", "Mixed-use", "Industri"],
    "sv": ["Bostad", "Kontor", "Logistik", "Hotell", "Mixed-use", "Industri"],
    "da": ["Bolig", "Kontor", "Logistik", "Hotel", "Mixed-use", "Industri"],
    "en-GB": ["Residential", "Office", "Logistics", "Hotel", "Mixed-use", "Industrial"],
    "en-US": ["Residential", "Office", "Logistics", "Hotel", "Mixed-use", "Industrial"],
    "de": ["Wohnen", "Büro", "Logistik", "Hotel", "Mixed-use", "Industrie"],
    "fi": ["Asuminen", "Toimisto", "Logistiikka", "Hotelli", "Mixed-use", "Teollisuus"],
}

LANG_LABELS = {
    "Norsk": "no",
    "Svenska": "sv",
    "Dansk": "da",
    "English (UK)": "en-GB",
    "English (US)": "en-US",
    "Deutsch": "de",
    "Suomi": "fi",
}

MARKETS: Dict[str, Dict[str, Any]] = {
    "NO": {
        "label": "Norge",
        "currency": "NOK",
        "vat": "25%",
        "rule_set": "TEK17 / EU Taxonomy / SFDR / ECB / Finanstilsynet",
        "jurisdiction_note": {
            "no": "Bruk TEK17 og norske offentlige datakilder som hovedramme. Kommunale krav og lokale planbestemmelser kan påvirke tolkning og tiltak.",
            "sv": "Använd TEK17 och norska offentliga datakällor som huvudram. Kommunala krav och lokala planbestämmelser kan påverka tolkning och åtgärder.",
            "da": "Brug TEK17 og norske offentlige datakilder som hovedramme. Kommunale krav og lokale planbestemmelser kan påvirke vurderingen.",
            "en-GB": "Use TEK17 and Norwegian public data as the baseline. Municipal requirements and local planning conditions may affect the final interpretation.",
            "en-US": "Use TEK17 and Norwegian public data as the baseline. Municipal requirements and local planning conditions may affect the final interpretation.",
            "de": "TEK17 und norwegische öffentliche Daten bilden die Grundlage. Kommunale Anforderungen und lokale Planvorgaben können die Bewertung beeinflussen.",
            "fi": "TEK17 ja norjalaiset julkiset tietolähteet muodostavat lähtökohdan. Kunnalliset vaatimukset ja paikalliset kaavamääräykset voivat vaikuttaa tulkintaan.",
        },
        "frameworks": [
            {"framework": "EU Taxonomy DNSH", "status": "ready", "format": "XML", "note": "Dokumenterer physical climate adaptation screening."},
            {"framework": "SFDR PAI", "status": "ready", "format": "CSV / JSON", "note": "Fysisk klimarisikoeksponering kan eksporteres videre."},
            {"framework": "ECB Climate Stress Test", "status": "ready", "format": "JSON", "note": "Felter for banker kan pre-populeres."},
            {"framework": "Finanstilsynet", "status": "ready", "format": "PDF", "note": "Norsk rapportstruktur for bank og forsikring."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "NVE Flomsonekart (WMS/WFS)", "access": "Public / registration"},
            {"hazard": "Landslide", "source": "NVE skred- og kvikkleirekart", "access": "Public / registration"},
            {"hazard": "Elevation", "source": "Kartverket DTM 1m", "access": "Indexed / local"},
            {"hazard": "Sea level", "source": "Kartverket / Bjerknessenteret", "access": "Periodic dataset"},
            {"hazard": "Heat", "source": "Copernicus Land Service", "access": "Public"},
            {"hazard": "Soil", "source": "NIBIO jordsmonnskart", "access": "Public"},
        ],
        "default_lang": "no",
    },
    "SE": {
        "label": "Sverige",
        "currency": "SEK",
        "vat": "25%",
        "rule_set": "PBL / Boverkets regler / EU Taxonomy / SFDR",
        "jurisdiction_note": {
            "sv": "Förklara när övergången 2025–2026 mellan äldre BBR och nyare Boverket-regler påverkar slutsatsen. Kommunala krav kan tillkomma.",
            "no": "Forklar når overgangen 2025–2026 mellom eldre BBR og nyere Boverket-regler påvirker konklusjonen. Kommunale krav kan komme i tillegg.",
            "da": "Forklar når overgangen 2025–2026 mellem ældre BBR og nyere Boverket-regler påvirker konklusionen. Kommunale krav kan komme oveni.",
            "en-GB": "Explain when the 2025–2026 transition between older BBR rules and newer Boverket regulations changes the conclusion. Municipal requirements may apply.",
            "en-US": "Explain when the 2025–2026 transition between older BBR rules and newer Boverket regulations changes the conclusion. Municipal requirements may apply.",
            "de": "Erklären Sie, wann der Übergang 2025–2026 zwischen älteren BBR-Regeln und neueren Boverket-Vorschriften die Schlussfolgerung beeinflusst. Kommunale Anforderungen können zusätzlich gelten.",
            "fi": "Selitä, milloin vuosien 2025–2026 siirtymä vanhojen BBR-sääntöjen ja uusien Boverket-määräysten välillä vaikuttaa johtopäätökseen. Kunnallisia vaatimuksia voi tulla lisää.",
        },
        "frameworks": [
            {"framework": "EU Taxonomy DNSH", "status": "ready", "format": "XML", "note": "No local adaptation required for core taxonomy logic."},
            {"framework": "SFDR PAI", "status": "ready", "format": "CSV / JSON", "note": "EU framework, same core fields as other EU markets."},
            {"framework": "ECB Climate Stress Test", "status": "ready", "format": "JSON", "note": "Bank portfolio screening compatible."},
            {"framework": "FI / Finansinspektionen", "status": "partial", "format": "PDF / JSON", "note": "Local supervisory wording should be confirmed per use case."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "MSB Översvämningsdatabas", "access": "Open WMS"},
            {"hazard": "Landslide", "source": "SGI rasriskkartor", "access": "Open"},
            {"hazard": "Elevation", "source": "Lantmäteriet höjdmodell", "access": "API / agreement"},
            {"hazard": "Energy", "source": "Boverket Energideklaration", "access": "Open API"},
        ],
        "default_lang": "sv",
    },
    "DK": {
        "label": "Danmark",
        "currency": "DKK",
        "vat": "25%",
        "rule_set": "BR18 / EU Taxonomy / SFDR",
        "jurisdiction_note": {
            "da": "Brug BR18 og danske datakilder som hovedramme. Kommunale og lokale forhold kan påvirke vurderingen.",
            "no": "Bruk BR18 og danske datakilder som hovedramme. Kommunale og lokale forhold kan påvirke vurderingen.",
            "sv": "Använd BR18 och danska datakällor som huvudram. Kommunala och lokala förhållanden kan påverka bedömningen.",
            "en-GB": "Use BR18 and Danish data sources as the baseline. Municipal and local conditions may affect the final assessment.",
            "en-US": "Use BR18 and Danish data sources as the baseline. Municipal and local conditions may affect the final assessment.",
            "de": "BR18 und dänische Datenquellen bilden die Grundlage. Kommunale und lokale Bedingungen können die Bewertung beeinflussen.",
            "fi": "BR18 ja tanskalaiset tietolähteet muodostavat lähtökohdan. Kunnalliset ja paikalliset olosuhteet voivat vaikuttaa arvioon.",
        },
        "frameworks": [
            {"framework": "EU Taxonomy DNSH", "status": "ready", "format": "XML", "note": "Same EU logic as Norway and Sweden."},
            {"framework": "SFDR PAI", "status": "ready", "format": "CSV / JSON", "note": "Physical risk export for funds and reporting."},
            {"framework": "ECB Climate Stress Test", "status": "ready", "format": "JSON", "note": "Bank screening compatible."},
            {"framework": "Finanstilsynet Danmark", "status": "partial", "format": "PDF / JSON", "note": "Local wording can be added in market templates."},
        ],
        "data_sources": [
            {"hazard": "Flood / sea", "source": "DMI Klimaatlas", "access": "Open API"},
            {"hazard": "Elevation", "source": "DHM / Kortforsyningen", "access": "Public services"},
            {"hazard": "Geology", "source": "GEUS", "access": "Public"},
            {"hazard": "Energy", "source": "Energistyrelsen Energimærke", "access": "API"},
        ],
        "default_lang": "da",
    },
    "FI": {
        "label": "Finland",
        "currency": "EUR",
        "vat": "24%",
        "rule_set": "RakMk / YM-asetukset / EU Taxonomy / SFDR",
        "jurisdiction_note": {
            "fi": "Suomessa julkisella sektorilla voi olla kaksikielisiä vaatimuksia (FI/SV). Käytä RakMk- ja YM-asetuksia lähtökohtana.",
            "no": "I Finland kan det være krav om tospråklighet (FI/SV), særlig i offentlig sektor. Bruk RakMk og YM-asetukset som grunnlag.",
            "sv": "I Finland kan tvåspråkiga krav (FI/SV) gälla, särskilt i offentlig sektor. Använd RakMk och YM-förordningar som utgångspunkt.",
            "da": "I Finland kan der gælde tosprogede krav (FI/SV), især i den offentlige sektor. Brug RakMk og YM-forordninger som udgangspunkt.",
            "en-GB": "Finland may require bilingual (FI/SV) outputs in parts of the public sector. Use RakMk and Ministry of the Environment regulations as the baseline.",
            "en-US": "Finland may require bilingual (FI/SV) outputs in parts of the public sector. Use RakMk and Ministry of the Environment regulations as the baseline.",
            "de": "In Finnland können zweisprachige Anforderungen (FI/SV) gelten, insbesondere im öffentlichen Sektor. RakMk und die Verordnungen des Umweltministeriums bilden die Grundlage.",
        },
        "frameworks": [
            {"framework": "EU Taxonomy DNSH", "status": "ready", "format": "XML", "note": "Same EU logic as other EU markets."},
            {"framework": "SFDR PAI", "status": "ready", "format": "CSV / JSON", "note": "Same EU logic as other EU markets."},
            {"framework": "ECB Climate Stress Test", "status": "ready", "format": "JSON", "note": "Portfolio screening for banks."},
            {"framework": "FIN-FSA / local reporting", "status": "partial", "format": "PDF / JSON", "note": "Local supervisory wording should be reviewed per use case."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "SYKE tulvakartat", "access": "Open WMS"},
            {"hazard": "Property", "source": "KTJ / Maanmittauslaitos", "access": "API"},
            {"hazard": "Energy", "source": "ARA energiatodistus", "access": "Search / API evolving"},
        ],
        "default_lang": "fi",
    },
    "UK": {
        "label": "United Kingdom (England default)",
        "currency": "GBP",
        "vat": "20%",
        "rule_set": "Building Regulations / Building Safety Act 2022 / FCA SDR",
        "jurisdiction_note": {
            "en-GB": "Default to England. Flag where Scotland, Wales or Northern Ireland may differ. Building Safety Act 2022 and post-Grenfell requirements can materially affect interpretation.",
            "en-US": "Default to England. Flag where Scotland, Wales, or Northern Ireland may differ. Building Safety Act 2022 and post-Grenfell requirements can materially affect interpretation.",
            "no": "Modulen bruker England som standard og flagger der Skottland, Wales eller Nord-Irland kan avvike. Building Safety Act 2022 etter Grenfell kan påvirke tolkningen vesentlig.",
            "sv": "Modulen använder England som standard och flaggar när Skottland, Wales eller Nordirland kan avvika. Building Safety Act 2022 kan påverka slutsatsen väsentligt.",
            "da": "Modulet bruger England som standard og markerer, hvor Skotland, Wales eller Nordirland kan afvige. Building Safety Act 2022 kan påvirke konklusionen væsentligt.",
            "de": "Das Modul verwendet England als Standard und kennzeichnet, wenn Schottland, Wales oder Nordirland abweichen können. Der Building Safety Act 2022 kann die Bewertung wesentlich beeinflussen.",
            "fi": "Moduuli käyttää Englantia oletuksena ja merkitsee tilanteet, joissa Skotlanti, Wales tai Pohjois-Irlanti voivat poiketa. Building Safety Act 2022 voi vaikuttaa arvioon olennaisesti.",
        },
        "frameworks": [
            {"framework": "UK Green Taxonomy", "status": "partial", "format": "XML / note", "note": "Still under development; not a direct EU Taxonomy mirror."},
            {"framework": "FCA SDR", "status": "ready", "format": "CSV / JSON", "note": "UK equivalent to SFDR-style sustainability disclosure."},
            {"framework": "TCFD / PRA", "status": "ready", "format": "PDF / JSON", "note": "Large-company and financial-sector reporting context."},
            {"framework": "Flood / planning", "status": "ready", "format": "PDF", "note": "Environment Agency flood planning outputs supported."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "Environment Agency Flood Map for Planning", "access": "Open API"},
            {"hazard": "Elevation", "source": "OS Terrain 5 / 50", "access": "OS Data Hub"},
            {"hazard": "Energy", "source": "EPC Open Data Communities", "access": "Open API"},
        ],
        "default_lang": "en-GB",
    },
    "DE": {
        "label": "Deutschland",
        "currency": "EUR",
        "vat": "19%",
        "rule_set": "Landesbauordnungen / BauGB / EU Taxonomy / SFDR",
        "jurisdiction_note": {
            "de": "Die anwendbare Landesbauordnung und die Praxis der zuständigen Behörde müssen bestätigt werden. Bayern oder NRW sind sinnvolle Pilotmärkte.",
            "no": "Relevant Landesbauordnung og lokal myndighetspraksis må bekreftes. Bayern eller NRW er naturlige pilotmarkeder.",
            "sv": "Tillämplig Landesbauordnung och lokal myndighetspraxis måste bekräftas. Bayern eller NRW är naturliga pilotmarknader.",
            "da": "Relevant Landesbauordnung og lokal myndighedspraksis skal bekræftes. Bayern eller NRW er naturlige pilotmarkeder.",
            "en-GB": "The applicable Landesbauordnung and local authority practice must be confirmed. Bavaria or NRW are sensible pilot states.",
            "en-US": "The applicable Landesbauordnung and local authority practice must be confirmed. Bavaria or NRW are sensible pilot states.",
            "fi": "Sovellettava Landesbauordnung ja paikallisen viranomaisen käytäntö on vahvistettava. Baijeri tai NRW ovat luontevia pilottialueita.",
        },
        "frameworks": [
            {"framework": "EU Taxonomy DNSH", "status": "ready", "format": "XML", "note": "Same core EU logic; local supervisory context via BaFin."},
            {"framework": "SFDR PAI", "status": "ready", "format": "CSV / JSON", "note": "Physical-risk outputs for EU fund reporting."},
            {"framework": "ECB Climate Stress Test", "status": "ready", "format": "JSON", "note": "Portfolio screening compatible."},
            {"framework": "BaFin context", "status": "partial", "format": "PDF / JSON", "note": "Local wording and state references should be reviewed."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "LAWA Hochwassergefahrenkarten", "access": "State portals / INSPIRE"},
            {"hazard": "Elevation", "source": "BKG DGM", "access": "Geobasisdaten services"},
            {"hazard": "Landslide", "source": "BGR geologische Gefahrenkarten", "access": "INSPIRE Geoportal"},
            {"hazard": "Property", "source": "ALKIS (state-based)", "access": "State integration required"},
        ],
        "default_lang": "de",
    },
    "US": {
        "label": "United States",
        "currency": "USD",
        "vat": "N/A",
        "rule_set": "State & local adoption / SEC / Fed-OCC guidance",
        "jurisdiction_note": {
            "en-US": "There is no single national building-code or property-data baseline. State and local adoption can materially change the result. Start with California or New York as pilot jurisdictions.",
            "en-GB": "There is no single national building-code or property-data baseline. State and local adoption can materially change the result. Start with California or New York as pilot jurisdictions.",
            "no": "Det finnes ingen én nasjonal byggeregel- eller eiendomsdatabase. Delstatlig og lokal kodeadopsjon kan endre vurderingen betydelig. Start med California eller New York som pilot.",
            "sv": "Det finns ingen enda nationell byggkod eller fastighetsdatabas. Delstatlig och lokal kodadoption kan ändra resultatet väsentligt. Börja med Kalifornien eller New York som pilot.",
            "da": "Der findes ikke én national byggelov eller ejendomsdatabase. Delstatslig og lokal kodeadoption kan ændre vurderingen væsentligt. Start med Californien eller New York som pilot.",
            "de": "Es gibt keine einheitliche nationale Bauordnung oder Eigentumsdatenbank. Staatliche und lokale Vorschriften können die Bewertung wesentlich verändern. California oder New York sind sinnvolle Pilotmärkte.",
            "fi": "Yhtä kansallista rakennusmääräystä tai kiinteistötietokantaa ei ole. Osavaltio- ja paikallissäännöt voivat muuttaa tulosta olennaisesti. Aloita Kaliforniasta tai New Yorkista pilottina.",
        },
        "frameworks": [
            {"framework": "SEC Climate Disclosure", "status": "partial", "format": "JSON / PDF", "note": "Federal climate-disclosure rules are evolving and partially challenged."},
            {"framework": "Federal Reserve / OCC", "status": "ready", "format": "JSON", "note": "Physical-risk fields for banking workflows."},
            {"framework": "FEMA / NFIP context", "status": "ready", "format": "PDF / JSON", "note": "Flood-zone classification tied to insurability and lending."},
            {"framework": "State climate rules", "status": "partial", "format": "PDF / note", "note": "California and New York differ materially from other states."},
        ],
        "data_sources": [
            {"hazard": "Flood", "source": "FEMA National Flood Hazard Layer (NFHL)", "access": "Open API"},
            {"hazard": "Elevation", "source": "USGS 3DEP", "access": "Open API"},
            {"hazard": "Landslide", "source": "USGS Landslide Hazard Program", "access": "Public / state supplements"},
            {"hazard": "Energy", "source": "ENERGY STAR / LEED / state sources", "access": "Mixed"},
        ],
        "default_lang": "en-US",
    },
}

SCENARIOS = ["RCP 4.5", "RCP 8.5"]
HORIZONS = ["2030", "2050", "2100"]
MARKET_OPTIONS = list(MARKETS.keys())
LANG_OPTIONS = list(LANG_LABELS.keys())

DISCLAIMER = {
    "no": "Dette er et AI-generert dataprodukt. Det er ikke en faglig attestasjon og kan ikke brukes som grunnlag for byggesøknad eller juridisk bindende avtale.",
    "sv": "Detta är en AI-genererad dataprodukt. Det är inte en fackmässig attestering och kan inte användas som underlag för bygglov eller juridiskt bindande avtal.",
    "da": "Dette er et AI-genereret dataprodukt. Det er ikke en faglig attestering og kan ikke bruges som grundlag for byggesag eller juridisk bindende aftaler.",
    "en-GB": "This is an AI-generated data product. It is not a professional attestation and cannot be used as the basis for statutory approval or legally binding commitments.",
    "en-US": "This is an AI-generated data product. It is not a professional attestation and cannot be used as the basis for statutory approval or legally binding commitments.",
    "de": "Dies ist ein KI-generiertes Datenprodukt. Es stellt keine fachliche Attestierung dar und kann nicht als Grundlage für behördliche Genehmigungen oder rechtlich bindende Verpflichtungen verwendet werden.",
    "fi": "Tämä on tekoälyn tuottama datapohjainen tuote. Se ei ole ammatillinen todistus eikä sitä voi käyttää viranomaisluvan tai oikeudellisesti sitovien sitoumusten perusteena.",
}

# ---------------------------------------------------------
# Helpers
# ---------------------------------------------------------

def tr(lang: str, key: str) -> str:
    return LANG[lang][key]


def safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value in (None, "", "-"):
            return default
        return float(value)
    except Exception:
        return default


def risk_band(score: float) -> str:
    if score >= 4.0:
        return "High"
    if score >= 3.0:
        return "Elevated"
    if score >= 2.0:
        return "Moderate"
    return "Low"


def localized_risk_band(score: float, lang: str) -> str:
    band = risk_band(score)
    labels = {
        "no": {"High": "Høy", "Elevated": "Forhøyet", "Moderate": "Moderat", "Low": "Lav"},
        "sv": {"High": "Hög", "Elevated": "Förhöjd", "Moderate": "Måttlig", "Low": "Låg"},
        "da": {"High": "Høj", "Elevated": "Forhøjet", "Moderate": "Moderat", "Low": "Lav"},
        "en-GB": {"High": "High", "Elevated": "Elevated", "Moderate": "Moderate", "Low": "Low"},
        "en-US": {"High": "High", "Elevated": "Elevated", "Moderate": "Moderate", "Low": "Low"},
        "de": {"High": "Hoch", "Elevated": "Erhöht", "Moderate": "Moderat", "Low": "Niedrig"},
        "fi": {"High": "Korkea", "Elevated": "Kohonnut", "Moderate": "Kohtalainen", "Low": "Matala"},
    }
    return labels.get(lang, labels["en-GB"])[band]


def format_money(value: float, currency: str) -> str:
    return f"{currency} {value:,.0f}".replace(",", " ")


def pick_soil_code(raw: str) -> str:
    raw = (raw or "").strip().lower()
    mapping = {
        "morene": "morene", "morän": "morene", "moræne": "morene", "moraine": "morene", "moräne": "morene", "moreeni": "morene",
        "leire": "clay", "lera": "clay", "ler": "clay", "clay": "clay", "lehm": "clay", "savi": "clay",
        "berg": "rock", "bedrock": "rock", "fels": "rock", "kallio": "rock",
        "marine avsetninger": "marine", "marine sediment": "marine", "marine sediments": "marine", "marina sediment": "marine", "marine aflejringer": "marine", "meri­sedimentti": "marine",
        "fyllmasser": "fill", "fyllnadsmassor": "fill", "fyldmasser": "fill", "made ground": "fill", "auffüllung": "fill", "täyttömaa": "fill",
    }
    return mapping.get(raw, "morene")


def soil_modifier(soil_code: str) -> float:
    return {"clay": 0.7, "marine": 0.8, "fill": 0.5, "morene": 0.1, "rock": -0.3}.get(soil_code, 0.0)


def market_multiplier(market: str) -> Dict[str, float]:
    # Subtle default calibration by market; internal calculations still follow shared model.
    if market == "DK":
        return {"flood": 1.05, "sea": 1.15, "landslide": 0.90, "heat": 1.00}
    if market == "SE":
        return {"flood": 1.00, "sea": 1.00, "landslide": 1.05, "heat": 1.00}
    if market == "FI":
        return {"flood": 1.00, "sea": 0.95, "landslide": 1.00, "heat": 0.95}
    if market == "UK":
        return {"flood": 1.10, "sea": 1.10, "landslide": 0.95, "heat": 1.00}
    if market == "DE":
        return {"flood": 1.05, "sea": 0.90, "landslide": 1.00, "heat": 1.00}
    if market == "US":
        return {"flood": 1.10, "sea": 1.05, "landslide": 1.00, "heat": 1.10}
    return {"flood": 1.0, "sea": 1.0, "landslide": 1.0, "heat": 1.0}


def compute_scores(asset: Dict[str, Any], market: str, weights: Dict[str, float]) -> Dict[str, Any]:
    elevation_m = safe_float(asset.get("elevation_m"), 12.0)
    distance_coast_km = safe_float(asset.get("distance_coast_km"), 2.0)
    distance_river_km = safe_float(asset.get("distance_river_km"), 0.8)
    slope_deg = safe_float(asset.get("slope_deg"), 6.0)
    heat_index = safe_float(asset.get("heat_index"), 5.0)
    basement = bool(asset.get("has_basement"))
    flood_zone = bool(asset.get("flood_zone"))
    landslide_zone = bool(asset.get("landslide_zone"))
    scenario = str(asset.get("scenario", "RCP 4.5"))
    horizon = str(asset.get("horizon", "2050"))
    asset_value = safe_float(asset.get("asset_value"), 50_000_000)
    soil_code = pick_soil_code(str(asset.get("soil_type", "Morene")))
    mult = market_multiplier(market)

    flood_score = 1.6 + (2.6 if flood_zone else 0.0) + max(0.0, 0.75 - distance_river_km) * 1.9
    flood_score += 0.25 if basement else 0.0
    flood_score -= elevation_m / 140.0
    flood_score *= mult["flood"]

    landslide_score = 1.4 + (2.4 if landslide_zone else 0.0) + slope_deg / 20.0 + soil_modifier(soil_code)
    landslide_score *= mult["landslide"]

    climate_modifier = 0.35 if "8.5" in scenario else 0.10
    horizon_modifier = 0.35 if horizon == "2100" else 0.18 if horizon == "2050" else 0.05

    sea_score = 4.15 - elevation_m / 21.0 - distance_coast_km / 2.0 + climate_modifier
    sea_score += 0.20 if basement else 0.0
    sea_score *= mult["sea"]

    heat_score = 1.3 + heat_index / 2.0 + horizon_modifier
    heat_score *= mult["heat"]

    flood_score = max(1.0, min(5.0, round(flood_score, 2)))
    landslide_score = max(1.0, min(5.0, round(landslide_score, 2)))
    sea_score = max(1.0, min(5.0, round(sea_score, 2)))
    heat_score = max(1.0, min(5.0, round(heat_score, 2)))

    total_weight = sum(weights.values()) or 1.0
    agg = (
        flood_score * weights["flood"]
        + landslide_score * weights["landslide"]
        + sea_score * weights["sea"]
        + heat_score * weights["heat"]
    ) / total_weight

    completeness_fields = [asset.get("address") or asset.get("lat"), asset.get("municipality") or asset.get("region"), asset.get("asset_class"), asset.get("asset_value")]
    completeness = sum(1 for v in completeness_fields if v not in (None, "", 0)) / len(completeness_fields)
    uncertainty = round(max(0.20, 1.10 - completeness * 0.35 - (0.15 if flood_zone else 0.0) - (0.10 if landslide_zone else 0.0)), 2)

    damage_cost = round(asset_value * (agg / 5.0) * (0.030 + (0.010 if flood_zone else 0.0) + (0.005 if basement else 0.0)), 0)
    top_factor = max(
        [("flood", flood_score), ("landslide", landslide_score), ("sea_level", sea_score), ("heat_stress", heat_score)],
        key=lambda x: x[1],
    )[0]

    return {
        "flood_score": flood_score,
        "landslide_score": landslide_score,
        "sea_level_score": sea_score,
        "heat_stress_score": heat_score,
        "aggregate_score": round(agg, 2),
        "uncertainty_interval": uncertainty,
        "estimated_damage_cost": damage_cost,
        "risk_band": risk_band(agg),
        "top_factor": top_factor,
    }


def top_factor_label(code: str, lang: str) -> str:
    labels = {
        "no": {"flood": "Flom", "landslide": "Skred / ras", "sea_level": "Havnivå / stormflo", "heat_stress": "Varmestress"},
        "sv": {"flood": "Översvämning", "landslide": "Ras / skred", "sea_level": "Havsnivå / stormflod", "heat_stress": "Värmestress"},
        "da": {"flood": "Oversvømmelse", "landslide": "Skred / ras", "sea_level": "Havniveau / stormflod", "heat_stress": "Varmestress"},
        "en-GB": {"flood": "Flood", "landslide": "Landslide", "sea_level": "Sea level / storm surge", "heat_stress": "Heat stress"},
        "en-US": {"flood": "Flood", "landslide": "Landslide", "sea_level": "Sea level / storm surge", "heat_stress": "Heat stress"},
        "de": {"flood": "Hochwasser", "landslide": "Rutschung / Hangversagen", "sea_level": "Meeresspiegel / Sturmflut", "heat_stress": "Hitzestress"},
        "fi": {"flood": "Tulva", "landslide": "Sortuma / maanvyöry", "sea_level": "Merenpinta / myrskytulva", "heat_stress": "Lämpöstressi"},
    }
    return labels.get(lang, labels["en-GB"]).get(code, code)


def make_source_rows(market: str, asset: Dict[str, Any]) -> List[Dict[str, str]]:
    rows = []
    for item in MARKETS[market]["data_sources"]:
        rows.append({
            "hazard": item["hazard"],
            "source": item["source"],
            "access": item["access"],
            "status": "configured" if asset.get("address") or asset.get("lat") else "waiting_for_input",
            "version": datetime.now(timezone.utc).strftime("%Y-%m"),
        })
    if asset.get("address") and market == "NO":
        geo = geocode_address(str(asset.get("address")), str(asset.get("municipality", "")))
        rows.append({
            "hazard": "Address resolution",
            "source": geo.get("source", "Address API"),
            "access": geo.get("status", "missing"),
            "status": geo.get("status", "missing"),
            "version": datetime.now(timezone.utc).strftime("%Y-%m"),
        })
    return rows


def make_regulatory_outputs(market: str, scores: Dict[str, Any], analysis_mode: str, asset_class: str) -> List[Dict[str, str]]:
    outputs = []
    for item in MARKETS[market]["frameworks"]:
        note = item["note"]
        if item["framework"].startswith("ECB") or item["framework"].startswith("Federal Reserve"):
            note = f"{note} Asset class: {asset_class}. Aggregate score: {scores['aggregate_score']}."
        outputs.append({
            "framework": item["framework"],
            "status": item["status"],
            "format": item["format"],
            "note": note,
        })
    if analysis_mode == "Portfolio":
        outputs.append({
            "framework": "Portfolio batch",
            "status": "ready",
            "format": "CSV / JSON",
            "note": "Batch summary with one row per asset and pre-populated risk fields.",
        })
    return outputs


def load_portfolio(upload) -> List[Dict[str, Any]]:
    if upload is None:
        return []
    name = getattr(upload, "name", "").lower()
    data = upload.getvalue()
    try:
        if name.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(data))
        else:
            df = pd.read_excel(io.BytesIO(data))
    except Exception:
        return []
    df = df.fillna("")
    rows: List[Dict[str, Any]] = []
    for _, row in df.head(5000).iterrows():
        rows.append({
            "id": row.get("id") or row.get("asset_id") or row.get("address") or f"asset-{len(rows)+1}",
            "asset_name": row.get("asset_name") or row.get("name") or row.get("property_name") or row.get("address") or f"Asset {len(rows)+1}",
            "address": row.get("address") or row.get("adresse") or "",
            "municipality": row.get("municipality") or row.get("kommune") or row.get("city") or "",
            "region": row.get("region") or row.get("state") or row.get("county") or "",
            "asset_class": row.get("asset_class") or row.get("type") or "Office",
            "asset_value": safe_float(row.get("asset_value") or row.get("value") or row.get("property_value"), 0.0),
            "elevation_m": safe_float(row.get("elevation_m") or row.get("elevation"), 12.0),
            "distance_coast_km": safe_float(row.get("distance_coast_km") or row.get("distance_to_coast_km"), 2.0),
            "distance_river_km": safe_float(row.get("distance_river_km") or row.get("distance_to_river_km"), 0.8),
            "slope_deg": safe_float(row.get("slope_deg") or row.get("slope"), 6.0),
            "heat_index": safe_float(row.get("heat_index") or row.get("urban_heat") or row.get("uhi_index"), 5.0),
            "soil_type": row.get("soil_type") or row.get("ground_type") or "Morene",
            "flood_zone": str(row.get("flood_zone") or row.get("in_flood_zone") or "").strip().lower() in {"1", "true", "yes", "ja", "y"},
            "landslide_zone": str(row.get("landslide_zone") or row.get("in_landslide_zone") or "").strip().lower() in {"1", "true", "yes", "ja", "y"},
            "has_basement": str(row.get("has_basement") or row.get("basement") or "").strip().lower() in {"1", "true", "yes", "ja", "y"},
        })
    return rows


def deterministic_summary(lang: str, market: str, asset_name: str, scores: Dict[str, Any], outputs: List[Dict[str, str]]) -> Dict[str, Any]:
    top = top_factor_label(scores["top_factor"], lang)
    band = localized_risk_band(scores["aggregate_score"], lang)
    framework_names = ", ".join(item["framework"] for item in outputs[:3])
    templates = {
        "no": f"{asset_name or 'Eiendommen'} har {band.lower()} klimarisiko med samlet score {scores['aggregate_score']}/5. Høyeste driver er {top.lower()}. Rapporten er strukturert for {framework_names} og kan brukes som screeninggrunnlag og videre oppfølging.",
        "sv": f"{asset_name or 'Fastigheten'} har {band.lower()} klimatrisk med total score {scores['aggregate_score']}/5. Starkaste riskdrivaren är {top.lower()}. Rapporten är strukturerad för {framework_names} och kan användas som screeningunderlag och vidare uppföljning.",
        "da": f"{asset_name or 'Ejendommen'} har {band.lower()} klimarisiko med samlet score {scores['aggregate_score']}/5. Den stærkeste risikodriver er {top.lower()}. Rapporten er struktureret til {framework_names} og kan bruges til screening og videre opfølgning.",
        "en-GB": f"{asset_name or 'The asset'} shows {band.lower()} climate risk with an aggregate score of {scores['aggregate_score']}/5. The main risk driver is {top.lower()}. The report is structured for {framework_names} and can be used for screening and follow-up.",
        "en-US": f"{asset_name or 'The asset'} shows {band.lower()} climate risk with an aggregate score of {scores['aggregate_score']}/5. The main risk driver is {top.lower()}. The report is structured for {framework_names} and can be used for screening and follow-up.",
        "de": f"{asset_name or 'Das Objekt'} weist ein {band.lower()}es Klimarisiko mit einem Gesamtscore von {scores['aggregate_score']}/5 auf. Der wichtigste Risikotreiber ist {top.lower()}. Der Bericht ist für {framework_names} strukturiert und eignet sich für Screening und weitere Maßnahmen.",
        "fi": f"{asset_name or 'Kohteessa'} on {band.lower()} ilmastoriski, kokonaispiste {scores['aggregate_score']}/5. Suurin riskiajuri on {top.lower()}. Raportti on jäsennelty kehyksille {framework_names} ja sitä voi käyttää seulontaan ja jatkotoimiin.",
    }
    action_templates = {
        "no": [
            {"action": f"Bekreft datagrunnlag for {top.lower()}", "priority": "High", "why": "Dette er faktoren som driver totalrisikoen mest."},
            {"action": "Avklar behov for tiltak eller dypere fagvurdering", "priority": "Medium", "why": "Høy eller forhøyet score bør følges opp før investering eller kredittbeslutning."},
            {"action": "Eksporter relevant regulatorisk mapping", "priority": "Medium", "why": "Gjør det enklere å bruke samme analyse i bank- og rapporteringsløp."},
        ],
        "sv": [
            {"action": f"Bekräfta datagrunden för {top.lower()}", "priority": "High", "why": "Detta är den faktor som driver totalrisken mest."},
            {"action": "Klargör behovet av åtgärder eller djupare fackgranskning", "priority": "Medium", "why": "Hög eller förhöjd score bör följas upp före investering eller kreditbeslut."},
            {"action": "Exportera relevant regulatorisk mapping", "priority": "Medium", "why": "Gör det enklare att använda samma analys i bank- och rapporteringsflöden."},
        ],
        "da": [
            {"action": f"Bekræft datagrundlaget for {top.lower()}", "priority": "High", "why": "Det er denne faktor, der driver totalrisikoen mest."},
            {"action": "Afklar behov for tiltag eller dybere faglig vurdering", "priority": "Medium", "why": "Høj eller forhøjet score bør følges op før investering eller kreditbeslutning."},
            {"action": "Eksportér relevant regulatorisk mapping", "priority": "Medium", "why": "Gør det lettere at bruge samme analyse i bank- og rapporteringsflows."},
        ],
        "en-GB": [
            {"action": f"Confirm the underlying data for {top.lower()}", "priority": "High", "why": "This factor is driving the total score the most."},
            {"action": "Clarify whether further mitigation or expert review is required", "priority": "Medium", "why": "High or elevated scores should be reviewed before investment or lending decisions."},
            {"action": "Export the relevant regulatory mapping", "priority": "Medium", "why": "This makes the same screening usable across banking and reporting workflows."},
        ],
        "en-US": [
            {"action": f"Confirm the underlying data for {top.lower()}", "priority": "High", "why": "This factor is driving the total score the most."},
            {"action": "Clarify whether further mitigation or expert review is required", "priority": "Medium", "why": "High or elevated scores should be reviewed before investment or lending decisions."},
            {"action": "Export the relevant regulatory mapping", "priority": "Medium", "why": "This makes the same screening usable across banking and reporting workflows."},
        ],
        "de": [
            {"action": f"Bestätigen Sie die Datengrundlage für {top.lower()}", "priority": "High", "why": "Dieser Faktor treibt den Gesamtscore am stärksten."},
            {"action": "Klären Sie, ob weitere Maßnahmen oder eine fachliche Prüfung erforderlich sind", "priority": "Medium", "why": "Hohe oder erhöhte Scores sollten vor Investitions- oder Kreditentscheidungen geprüft werden."},
            {"action": "Exportieren Sie das relevante regulatorische Mapping", "priority": "Medium", "why": "So kann dieselbe Analyse in Bank- und Reporting-Workflows genutzt werden."},
        ],
        "fi": [
            {"action": f"Vahvista {top.lower()}-riskin tietoperusta", "priority": "High", "why": "Tämä tekijä vaikuttaa kokonaispisteeseen eniten."},
            {"action": "Arvioi tarvitaanko lisätoimia tai asiantuntijakatselmusta", "priority": "Medium", "why": "Korkea tai kohonnut pistetaso kannattaa tarkistaa ennen sijoitus- tai luottopäätöstä."},
            {"action": "Vie tarvittava sääntelykartoitus", "priority": "Medium", "why": "Näin sama analyysi on käytettävissä pankki- ja raportointiprosesseissa."},
        ],
    }
    return {
        "executive_summary": templates.get(lang, templates["en-GB"]),
        "key_drivers": [top, band, f"{scores['aggregate_score']}/5"],
        "recommended_actions": action_templates.get(lang, action_templates["en-GB"]),
    }


def ai_summary(lang: str, market: str, analysis_payload: Dict[str, Any]) -> Dict[str, Any]:
    if generate_json_with_fallback is None:
        return {"ok": False, "error": "AI fallback unavailable", "attempt_log": []}

    schema = {
        "executive_summary": "Short end-user summary",
        "key_drivers": ["List the main drivers"],
        "recommended_actions": [{"action": "", "priority": "High|Medium|Low", "why": ""}],
        "jurisdiction_note": "Short country-specific legal note",
    }
    market_cfg = MARKETS[market]
    lang_name = {
        "no": "Norsk bokmål", "sv": "Svenska", "da": "Dansk", "en-GB": "British English", "en-US": "American English", "de": "Deutsch", "fi": "Suomi"
    }[lang]
    system_prompt = f"""
You are the Builtly Climate Risk report writer.
Write in {lang_name}.
Audience: end users such as developers, banks, insurers, investors and project teams.
Jurisdiction: {market_cfg['label']}.
Primary rule set / reporting context: {market_cfg['rule_set']}.
Country note: {market_cfg['jurisdiction_note'][lang]}.
This module is a level 1 auto-report product. Do not claim legal approval, certification or professional sign-off.
Be practical, concise, and user-facing.
Return JSON only.
""".strip()
    user_prompt = json.dumps(analysis_payload, ensure_ascii=False, indent=2)
    return generate_json_with_fallback(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        schema_hint=schema,
        task="document_engine",
        preferred_providers=["openai", "anthropic", "gemini"],
        estimated_context_chars=len(user_prompt),
        max_output_tokens=1000,
        temperature=0.15,
    )


def build_factor_df(lang: str, scores: Dict[str, Any], source_rows: List[Dict[str, str]]) -> pd.DataFrame:
    label_map = {
        "no": {"flood": "Flom", "landslide": "Skred / ras", "sea_level": "Havnivå / stormflo", "heat_stress": "Varmestress"},
        "sv": {"flood": "Översvämning", "landslide": "Ras / skred", "sea_level": "Havsnivå / stormflod", "heat_stress": "Värmestress"},
        "da": {"flood": "Oversvømmelse", "landslide": "Skred / ras", "sea_level": "Havniveau / stormflod", "heat_stress": "Varmestress"},
        "en-GB": {"flood": "Flood", "landslide": "Landslide", "sea_level": "Sea level / storm surge", "heat_stress": "Heat stress"},
        "en-US": {"flood": "Flood", "landslide": "Landslide", "sea_level": "Sea level / storm surge", "heat_stress": "Heat stress"},
        "de": {"flood": "Hochwasser", "landslide": "Rutschung", "sea_level": "Meeresspiegel", "heat_stress": "Hitzestress"},
        "fi": {"flood": "Tulva", "landslide": "Maanvyöry", "sea_level": "Merenpinta", "heat_stress": "Lämpöstressi"},
    }
    labels = label_map.get(lang, label_map["en-GB"])
    confidence = max(0.45, min(0.95, 1.0 - scores["uncertainty_interval"] / 2.0))
    source_text = ", ".join({row["source"] for row in source_rows[:3]})
    rows = [
        {"factor": labels["flood"], "score": scores["flood_score"], "confidence": round(confidence, 2), "note": "risk driver", "source": source_text},
        {"factor": labels["landslide"], "score": scores["landslide_score"], "confidence": round(confidence, 2), "note": "risk driver", "source": source_text},
        {"factor": labels["sea_level"], "score": scores["sea_level_score"], "confidence": round(confidence, 2), "note": "risk driver", "source": source_text},
        {"factor": labels["heat_stress"], "score": scores["heat_stress_score"], "confidence": round(confidence, 2), "note": "risk driver", "source": source_text},
    ]
    return pd.DataFrame(rows)


def portfolio_analysis(rows: List[Dict[str, Any]], market: str, scenario: str, horizon: str, weights: Dict[str, float]) -> pd.DataFrame:
    out_rows: List[Dict[str, Any]] = []
    for row in rows:
        asset = dict(row)
        asset["scenario"] = scenario
        asset["horizon"] = horizon
        scores = compute_scores(asset, market, weights)
        out_rows.append({
            "id": row.get("id") or row.get("asset_name"),
            "asset_name": row.get("asset_name") or row.get("address") or row.get("id"),
            "address": row.get("address", ""),
            "municipality": row.get("municipality", ""),
            "region": row.get("region", ""),
            "asset_class": row.get("asset_class", ""),
            "aggregate_score": scores["aggregate_score"],
            "risk_band": scores["risk_band"],
            "top_factor": scores["top_factor"],
            "estimated_damage_cost": scores["estimated_damage_cost"],
            "uncertainty_interval": scores["uncertainty_interval"],
        })
    return pd.DataFrame(out_rows)


def _bytes_io() -> io.BytesIO:
    return io.BytesIO()


def xml_from_mapping(meta: Dict[str, Any], regulatory_outputs: List[Dict[str, str]], factor_df: pd.DataFrame) -> bytes:
    root = ET.Element("climateRiskAnalysis")
    meta_el = ET.SubElement(root, "meta")
    for key, value in meta.items():
        child = ET.SubElement(meta_el, key)
        child.text = str(value)
    factors_el = ET.SubElement(root, "riskFactors")
    for _, row in factor_df.iterrows():
        factor = ET.SubElement(factors_el, "factor")
        for key in ["factor", "score", "confidence", "source"]:
            child = ET.SubElement(factor, key)
            child.text = str(row[key])
    reg_el = ET.SubElement(root, "regulatoryOutputs")
    for item in regulatory_outputs:
        out = ET.SubElement(reg_el, "output")
        for key in ["framework", "status", "format", "note"]:
            child = ET.SubElement(out, key)
            child.text = str(item.get(key, ""))
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def markdown_report(lang: str, market: str, meta: Dict[str, Any], summary: str, factor_df: pd.DataFrame, regulatory_outputs: List[Dict[str, str]], source_rows: List[Dict[str, str]], actions: List[Dict[str, str]]) -> str:
    lines = [
        f"# {meta['report_title']}",
        "",
        f"- Analysis ID: {meta['analysis_id']}",
        f"- Market: {MARKETS[market]['label']}",
        f"- Asset: {meta['asset_name']}",
        f"- Created: {meta['created_at']}",
        f"- Delivery level: auto",
        "",
        "## Summary",
        summary,
        "",
        "## Scores",
    ]
    for _, row in factor_df.iterrows():
        lines.append(f"- **{row['factor']}**: {row['score']} (confidence {row['confidence']})")
    lines += ["", "## Regulatory mapping"]
    for item in regulatory_outputs:
        lines.append(f"- **{item['framework']}** ({item['status']}): {item['note']}")
    lines += ["", "## Data sources"]
    for row in source_rows:
        lines.append(f"- **{row['source']}** – {row['hazard']} ({row['access']})")
    lines += ["", "## Recommended actions"]
    for action in actions:
        lines.append(f"- **{action['priority']}**: {action['action']} — {action['why']}")
    lines += ["", "## Disclaimer", DISCLAIMER[lang], ""]
    return "\n".join(lines)


def build_docx_report(lang: str, market: str, meta: Dict[str, Any], summary: str, factor_df: pd.DataFrame, regulatory_outputs: List[Dict[str, str]], source_rows: List[Dict[str, str]], actions: List[Dict[str, str]]) -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    title = doc.add_heading(meta["report_title"], 0)
    if WD_ALIGN_PARAGRAPH:
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    p.add_run(f"Analysis ID: {meta['analysis_id']}\n").bold = True
    p.add_run(f"Created: {meta['created_at']}\n")
    p.add_run(f"Market: {MARKETS[market]['label']}\n")
    p.add_run(f"Asset: {meta['asset_name']}\n")
    p.add_run(f"Delivery level: auto")

    doc.add_heading(tr(lang, "summary_title"), level=1)
    doc.add_paragraph(summary)

    doc.add_heading(tr(lang, "factors_title"), level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = tr(lang, "tab_factors")
    hdr[1].text = "Score"
    hdr[2].text = "Confidence"
    hdr[3].text = "Source"
    for _, row in factor_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["factor"])
        cells[1].text = str(row["score"])
        cells[2].text = str(row["confidence"])
        cells[3].text = str(row["source"])

    doc.add_heading(tr(lang, "mapping_title"), level=1)
    for item in regulatory_outputs:
        doc.add_paragraph(f"{item['framework']} ({item['status']}) — {item['note']}", style="List Bullet")

    doc.add_heading(tr(lang, "sources_title"), level=1)
    for row in source_rows:
        doc.add_paragraph(f"{row['source']} — {row['hazard']} — {row['access']}", style="List Bullet")

    doc.add_heading(tr(lang, "actions_title"), level=1)
    for action in actions:
        doc.add_paragraph(f"{action['priority']}: {action['action']} — {action['why']}", style="List Bullet")

    doc.add_heading(tr(lang, "disclaimer_title"), level=1)
    doc.add_paragraph(DISCLAIMER[lang])

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_pdf_report(lang: str, market: str, meta: Dict[str, Any], summary: str, factor_df: pd.DataFrame, regulatory_outputs: List[Dict[str, str]], source_rows: List[Dict[str, str]], actions: List[Dict[str, str]]) -> Optional[bytes]:
    if not REPORTLAB_READY:
        return None
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)

    try:
        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
        if Path(font_path).exists() and Path(bold_path).exists():
            pdfmetrics.registerFont(TTFont("BuiltlySans", font_path))
            pdfmetrics.registerFont(TTFont("BuiltlySansBold", bold_path))
            base_font = "BuiltlySans"
            bold_font = "BuiltlySansBold"
        else:
            base_font = "Helvetica"
            bold_font = "Helvetica-Bold"
    except Exception:
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BuiltlyTitle", fontName=bold_font, fontSize=18, leading=22, textColor=colors.HexColor("#0b1f3a")))
    styles.add(ParagraphStyle(name="BuiltlyHead", fontName=bold_font, fontSize=12, leading=16, textColor=colors.HexColor("#18325a"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="BuiltlyBody", fontName=base_font, fontSize=9.6, leading=14))

    story = [
        Paragraph(meta["report_title"], styles["BuiltlyTitle"]),
        Spacer(1, 4 * mm),
        Paragraph(f"Analysis ID: {meta['analysis_id']}<br/>Created: {meta['created_at']}<br/>Market: {MARKETS[market]['label']}<br/>Asset: {meta['asset_name']}<br/>Delivery level: auto", styles["BuiltlyBody"]),
        Spacer(1, 4 * mm),
        Paragraph(tr(lang, "summary_title"), styles["BuiltlyHead"]),
        Paragraph(summary, styles["BuiltlyBody"]),
        Spacer(1, 3 * mm),
        Paragraph(tr(lang, "factors_title"), styles["BuiltlyHead"]),
    ]
    factor_table = Table([[str(c) for c in factor_df.columns]] + factor_df.astype(str).values.tolist(), repeatRows=1)
    factor_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#19324d")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), bold_font),
        ("FONTNAME", (0, 1), (-1, -1), base_font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d3d9e4")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#eef3f8")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += [factor_table, Spacer(1, 3 * mm), Paragraph(tr(lang, "mapping_title"), styles["BuiltlyHead"])]
    for item in regulatory_outputs:
        story.append(Paragraph(f"<b>{item['framework']}</b> ({item['status']}) – {item['note']}", styles["BuiltlyBody"]))
    story += [Spacer(1, 3 * mm), Paragraph(tr(lang, "sources_title"), styles["BuiltlyHead"])]
    for row in source_rows:
        story.append(Paragraph(f"{row['source']} – {row['hazard']} – {row['access']}", styles["BuiltlyBody"]))
    story += [Spacer(1, 3 * mm), Paragraph(tr(lang, "actions_title"), styles["BuiltlyHead"])]
    for action in actions:
        story.append(Paragraph(f"<b>{action['priority']}</b>: {action['action']} – {action['why']}", styles["BuiltlyBody"]))
    story += [Spacer(1, 3 * mm), Paragraph(tr(lang, "disclaimer_title"), styles["BuiltlyHead"]), Paragraph(DISCLAIMER[lang], styles["BuiltlyBody"])]
    doc.build(story)
    return bio.getvalue()


def audit_entry(provider_result: Dict[str, Any], analysis_id: str, market: str, lang: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "analysis_id": analysis_id,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "delivery_level": "auto",
        "market": market,
        "language": lang,
        "provider": provider_result.get("provider") or "deterministic",
        "model": provider_result.get("model") or "deterministic-template",
        "attempt_log": provider_result.get("attempt_log", []),
        "input_summary": {
            "analysis_mode": payload.get("analysis_mode"),
            "asset_count": payload.get("asset_count"),
            "scenario": payload.get("scenario"),
            "horizon": payload.get("horizon"),
        },
    }


# ---------------------------------------------------------
# Main page
# ---------------------------------------------------------

project = configure_page("Builtly | Climate Risk", "🌍")

# Defaults from project state
market_default = "NO"
market_names = {v["label"]: k for k, v in MARKETS.items()}
lang_default_label = "Norsk"
if str(project.get("land", "")).lower().startswith("sverige"):
    market_default, lang_default_label = "SE", "Svenska"
elif str(project.get("land", "")).lower().startswith("danmark"):
    market_default, lang_default_label = "DK", "Dansk"
elif "united states" in str(project.get("land", "")).lower():
    market_default, lang_default_label = "US", "English (US)"
elif "united kingdom" in str(project.get("land", "")).lower() or "england" in str(project.get("land", "")).lower():
    market_default, lang_default_label = "UK", "English (UK)"
elif "tysk" in str(project.get("land", "")).lower() or "deutsch" in str(project.get("land", "")).lower():
    market_default, lang_default_label = "DE", "Deutsch"
elif "finland" in str(project.get("land", "")).lower() or "suomi" in str(project.get("land", "")).lower():
    market_default, lang_default_label = "FI", "Suomi"

lang_label = st.session_state.get("climate_lang_label", lang_default_label)
if lang_label not in LANG_LABELS:
    lang_label = lang_default_label
lang = LANG_LABELS[lang_label]

render_hero(
    eyebrow=tr(lang, "eyebrow"),
    title=tr(lang, "title"),
    subtitle=tr(lang, "subtitle"),
    pills=["Flood", "Landslide", "Sea level", "Heat stress", "Portfolio", "PDF / DOCX / JSON / XML"],
    badge=tr(lang, "eyebrow"),
)

render_section(tr(lang, "setup_title"), tr(lang, "setup_sub"), tr(lang, "setup_kicker"))
left, right = st.columns([1.25, 0.75], gap="large")

with left:
    lang_idx = LANG_OPTIONS.index(lang_label)
    lang_label = st.selectbox(tr(lang, "lang_label"), LANG_OPTIONS, index=lang_idx)
    lang = LANG_LABELS[lang_label]
    st.session_state["climate_lang_label"] = lang_label

    market_labels = [f"{MARKETS[k]['label']} ({k})" for k in MARKET_OPTIONS]
    market_idx = MARKET_OPTIONS.index(st.session_state.get("climate_market", market_default)) if st.session_state.get("climate_market", market_default) in MARKET_OPTIONS else MARKET_OPTIONS.index(market_default)
    market_choice = st.selectbox(tr(lang, "market_label"), market_labels, index=market_idx)
    market = market_choice.rsplit("(", 1)[1].replace(")", "")
    st.session_state["climate_market"] = market
    market_cfg = MARKETS[market]

    mode = st.radio(tr(lang, "mode_label"), [tr(lang, "mode_single"), tr(lang, "mode_portfolio")], horizontal=True)
    analysis_mode = "Portfolio" if mode == tr(lang, "mode_portfolio") else "Single"

    a1, a2 = st.columns(2)
    with a1:
        asset_name = st.text_input(tr(lang, "property_name"), value=project.get("p_name", ""))
        address = st.text_input(tr(lang, "address"), value=project.get("adresse", ""))
        municipality = st.text_input(tr(lang, "municipality"), value=project.get("kommune", ""))
        region = st.text_input(tr(lang, "region"), value="")
        asset_class = st.selectbox(tr(lang, "asset_class"), ASSET_CLASSES.get(lang, ASSET_CLASSES["en-GB"]), index=1)
    with a2:
        year_built = st.number_input(tr(lang, "year_built"), min_value=1850, max_value=2100, value=2012, step=1)
        area_m2 = st.number_input(tr(lang, "area"), min_value=50.0, value=float(project.get("bta", 2500)), step=50.0)
        currency = market_cfg["currency"]
        st.text_input(tr(lang, "currency_hint"), value=currency, disabled=True)
        asset_value = st.number_input(tr(lang, "value"), min_value=0.0, value=85_000_000.0 if currency in {"NOK", "SEK", "DKK"} else 12_000_000.0, step=100_000.0)

    st.markdown(f"### {tr(lang, 'analysis_settings')}")
    b1, b2, b3, b4 = st.columns(4)
    with b1:
        scenario = st.selectbox(tr(lang, "scenario"), SCENARIOS, index=0)
        horizon = st.selectbox(tr(lang, "horizon"), HORIZONS, index=1)
    with b2:
        elevation_m = st.number_input(tr(lang, "elevation"), min_value=0.0, value=14.0, step=1.0)
        distance_coast_km = st.number_input(tr(lang, "coast"), min_value=0.0, value=1.8, step=0.1)
    with b3:
        distance_river_km = st.number_input(tr(lang, "river"), min_value=0.0, value=0.45, step=0.05)
        slope_deg = st.number_input(tr(lang, "slope"), min_value=0.0, value=12.0, step=1.0)
    with b4:
        soil_type = st.selectbox(tr(lang, "soil"), SOIL_OPTIONS.get(lang, SOIL_OPTIONS["en-GB"]), index=1)
        heat_index = st.slider(tr(lang, "heat"), 0, 10, 6)

    c1, c2, c3 = st.columns(3)
    with c1:
        flood_zone = st.toggle(tr(lang, "flood_zone"), value=True)
    with c2:
        landslide_zone = st.toggle(tr(lang, "landslide_zone"), value=False)
    with c3:
        has_basement = st.toggle(tr(lang, "basement"), value=True)

    st.markdown(f"### {tr(lang, 'weights_title')}")
    st.caption(tr(lang, "weights_help"))
    w1, w2, w3, w4 = st.columns(4)
    with w1:
        weight_flood = st.slider("Flood", 0.0, 1.0, 0.35, 0.05)
    with w2:
        weight_landslide = st.slider("Landslide", 0.0, 1.0, 0.25, 0.05)
    with w3:
        weight_sea = st.slider("Sea", 0.0, 1.0, 0.25, 0.05)
    with w4:
        weight_heat = st.slider("Heat", 0.0, 1.0, 0.15, 0.05)
    weights = {"flood": weight_flood, "landslide": weight_landslide, "sea": weight_sea, "heat": weight_heat}

    portfolio_upload = None
    if analysis_mode == "Portfolio":
        portfolio_upload = st.file_uploader(tr(lang, "upload_portfolio"), type=["csv", "xlsx", "xls"], key="climate_portfolio_upload_v2")

    st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
    can_run = bool(address or portfolio_upload or asset_name)
    generate_clicked = st.button(tr(lang, "generate"), type="primary", use_container_width=True)
    rerun_clicked = st.button(tr(lang, "rerun"), use_container_width=True)
    if (generate_clicked or rerun_clicked) and not can_run:
        st.warning(tr(lang, "need_inputs"))

with right:
    market_cfg = MARKETS[st.session_state.get("climate_market", market_default)]
    render_project_snapshot(project, badge=market_cfg["label"])
    render_panel(
        tr(lang, "live_data_title"),
        tr(lang, "live_data_sub"),
        [
            f"{tr(lang, 'country_note_label')}: {market_cfg['jurisdiction_note'][lang]}",
            f"Rule set: {market_cfg['rule_set']}",
            f"Currency: {market_cfg['currency']} · VAT: {market_cfg['vat']}",
        ],
        tone="blue",
        badge=market_cfg["label"],
    )
    render_panel(
        tr(lang, "disclaimer_title"),
        DISCLAIMER[lang],
        [],
        tone="gold",
        badge="Level 1",
    )

# Analysis execution
if generate_clicked or rerun_clicked:
    market = st.session_state.get("climate_market", market_default)
    market_cfg = MARKETS[market]
    asset_payload = {
        "asset_name": asset_name or project.get("p_name") or "Asset",
        "address": address,
        "municipality": municipality,
        "region": region,
        "asset_class": asset_class,
        "year_built": year_built,
        "area_m2": area_m2,
        "asset_value": asset_value,
        "elevation_m": elevation_m,
        "distance_coast_km": distance_coast_km,
        "distance_river_km": distance_river_km,
        "slope_deg": slope_deg,
        "soil_type": soil_type,
        "heat_index": heat_index,
        "flood_zone": flood_zone,
        "landslide_zone": landslide_zone,
        "has_basement": has_basement,
        "scenario": scenario,
        "horizon": horizon,
    }
    single_scores = compute_scores(asset_payload, market, weights)
    source_rows = make_source_rows(market, asset_payload)
    factor_df = build_factor_df(lang, single_scores, source_rows)
    reg_outputs = make_regulatory_outputs(market, single_scores, "Portfolio" if analysis_mode == "Portfolio" else "Single", asset_class)

    portfolio_df = None
    portfolio_rows = []
    portfolio_summary = {}
    if analysis_mode == "Portfolio":
        portfolio_rows = load_portfolio(portfolio_upload)
        if not portfolio_rows:
            portfolio_rows = [
                dict(asset_payload, id=f"asset-{i+1}", asset_name=f"Asset {i+1}")
                for i in range(1, 11)
            ]
        portfolio_df = portfolio_analysis(portfolio_rows, market, scenario, horizon, weights)
        if not portfolio_df.empty:
            portfolio_summary = {
                "count": int(len(portfolio_df)),
                "avg_score": round(float(portfolio_df["aggregate_score"].mean()), 2),
                "high_count": int((portfolio_df["aggregate_score"] >= 4.0).sum()),
                "elevated_count": int(((portfolio_df["aggregate_score"] >= 3.0) & (portfolio_df["aggregate_score"] < 4.0)).sum()),
                "total_damage": float(portfolio_df["estimated_damage_cost"].sum()),
            }

    analysis_id = f"cr-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
    ai_payload = {
        "analysis_mode": analysis_mode,
        "asset_count": 1 if analysis_mode == "Single" else len(portfolio_rows),
        "market": market_cfg["label"],
        "rule_set": market_cfg["rule_set"],
        "scenario": scenario,
        "horizon": horizon,
        "scores": single_scores,
        "top_factor": top_factor_label(single_scores["top_factor"], lang),
        "regulatory_outputs": reg_outputs,
        "portfolio_summary": portfolio_summary,
    }
    ai_result = ai_summary(lang, market, ai_payload)
    if ai_result.get("ok") and ai_result.get("data"):
        ai_data = ai_result["data"]
        summary_text = ai_data.get("executive_summary") or deterministic_summary(lang, market, asset_payload["asset_name"], single_scores, reg_outputs)["executive_summary"]
        actions = ai_data.get("recommended_actions") or deterministic_summary(lang, market, asset_payload["asset_name"], single_scores, reg_outputs)["recommended_actions"]
        jurisdiction_note = ai_data.get("jurisdiction_note") or market_cfg["jurisdiction_note"][lang]
    else:
        fallback = deterministic_summary(lang, market, asset_payload["asset_name"], single_scores, reg_outputs)
        summary_text = fallback["executive_summary"]
        actions = fallback["recommended_actions"]
        jurisdiction_note = market_cfg["jurisdiction_note"][lang]

    meta = {
        "analysis_id": analysis_id,
        "asset_name": asset_payload["asset_name"],
        "report_title": f"Builtly Climate Risk Report – {asset_payload['asset_name']}" if lang.startswith("en") else (
            f"Builtly Klimarisikorapport – {asset_payload['asset_name']}" if lang in {"no", "da"} else
            f"Builtly Klimatriskrapport – {asset_payload['asset_name']}" if lang == "sv" else
            f"Builtly Klimarisikobericht – {asset_payload['asset_name']}" if lang == "de" else
            f"Builtly Ilmastoriskiraportti – {asset_payload['asset_name']}"
        ),
        "created_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
    }

    package_json = {
        "meta": {**meta, "market": market, "market_label": market_cfg["label"], "delivery_level": "auto", "language": lang, "jurisdiction_note": jurisdiction_note},
        "input": asset_payload,
        "scores": single_scores,
        "regulatory_outputs": reg_outputs,
        "source_rows": source_rows,
        "summary": summary_text,
        "actions": actions,
        "portfolio_summary": portfolio_summary,
        "portfolio_rows": portfolio_df.to_dict(orient="records") if portfolio_df is not None else [],
    }
    md_report = markdown_report(lang, market, meta, summary_text, factor_df, reg_outputs, source_rows, actions)
    docx_bytes = build_docx_report(lang, market, meta, summary_text, factor_df, reg_outputs, source_rows, actions)
    pdf_bytes = build_pdf_report(lang, market, meta, summary_text, factor_df, reg_outputs, source_rows, actions)
    xml_bytes = xml_from_mapping({
        "analysis_id": analysis_id,
        "market": market,
        "asset_name": asset_payload["asset_name"],
        "aggregate_score": single_scores["aggregate_score"],
        "created_at": meta["created_at"],
    }, reg_outputs, factor_df)
    audit = audit_entry(ai_result, analysis_id, market, lang, ai_payload)

    st.session_state["climate_report_state"] = {
        "lang": lang,
        "market": market,
        "asset_payload": asset_payload,
        "scores": single_scores,
        "factor_df": factor_df,
        "reg_outputs": reg_outputs,
        "source_rows": source_rows,
        "summary": summary_text,
        "actions": actions,
        "jurisdiction_note": jurisdiction_note,
        "portfolio_df": portfolio_df,
        "portfolio_summary": portfolio_summary,
        "package_json": package_json,
        "pdf_bytes": pdf_bytes,
        "docx_bytes": docx_bytes,
        "xml_bytes": xml_bytes,
        "md_report": md_report,
        "audit": audit,
        "attempt_log": ai_result.get("attempt_log", []),
        "ai_ok": ai_result.get("ok", False),
        "analysis_id": analysis_id,
    }

state = st.session_state.get("climate_report_state")
if state:
    lang = state["lang"]
    market = state["market"]
    market_cfg = MARKETS[market]
    scores = state["scores"]
    factor_df = state["factor_df"]
    reg_outputs = state["reg_outputs"]
    source_rows = state["source_rows"]
    portfolio_df = state["portfolio_df"]
    portfolio_summary = state["portfolio_summary"]
    render_section(tr(lang, "results_title"), tr(lang, "results_sub"), tr(lang, "results_kicker"))
    framework_primary = reg_outputs[0]["framework"] if reg_outputs else market_cfg["rule_set"]
    render_metric_cards([
        {"label": tr(lang, "metrics_score"), "value": f"{scores['aggregate_score']}/5", "desc": f"{tr(lang, 'risk_band')}: {localized_risk_band(scores['aggregate_score'], lang)}"},
        {"label": tr(lang, "metrics_uncertainty"), "value": f"± {scores['uncertainty_interval']}", "desc": f"{tr(lang, 'top_risk')}: {top_factor_label(scores['top_factor'], lang)}"},
        {"label": tr(lang, "metrics_assets"), "value": str(1 if portfolio_df is None else len(portfolio_df)), "desc": market_cfg['label']},
        {"label": tr(lang, "metrics_framework"), "value": framework_primary, "desc": market_cfg['rule_set']},
    ])

    tabs = st.tabs([
        tr(lang, "tab_overview"),
        tr(lang, "tab_factors"),
        tr(lang, "tab_mapping"),
        tr(lang, "tab_sources"),
        tr(lang, "tab_package"),
        tr(lang, "tab_portfolio"),
    ])

    with tabs[0]:
        render_panel(
            tr(lang, "summary_title"),
            state["summary"],
            [
                f"{tr(lang, 'country_note_label')}: {state['jurisdiction_note']}",
                tr(lang, "map_hint"),
            ],
            tone=tone_from_score(scores["aggregate_score"]),
            badge=localized_risk_band(scores["aggregate_score"], lang),
        )
        render_panel(
            tr(lang, "actions_title"),
            "",
            [f"{item['priority']}: {item['action']} — {item['why']}" for item in state["actions"]],
            tone="green",
            badge="Action",
        )
        if not state["ai_ok"]:
            st.info(tr(lang, "ai_unavailable"))

    with tabs[1]:
        st.bar_chart(factor_df.set_index("factor")["score"])
        st.dataframe(factor_df, use_container_width=True, hide_index=True)
        dataframe_download(factor_df, tr(lang, "download_csv"), f"climate_factors_{state['analysis_id']}.csv")

    with tabs[2]:
        reg_df = pd.DataFrame(reg_outputs)
        st.dataframe(reg_df, use_container_width=True, hide_index=True)
        st.caption(state["jurisdiction_note"])

    with tabs[3]:
        source_df = pd.DataFrame(source_rows)
        st.dataframe(source_df, use_container_width=True, hide_index=True)
        render_panel(
            tr(lang, "live_data_title"),
            tr(lang, "live_data_sub"),
            [f"{row['source']} — {row['hazard']} — {row['access']}" for row in source_rows],
            tone="blue",
            badge=market_cfg["label"],
        )

    with tabs[4]:
        st.success(tr(lang, "analysis_ready"))
        st.caption(tr(lang, "report_sub"))
        disclaimer_accepted = st.checkbox(tr(lang, "disclaimer_ack"), key=f"climate_disclaimer_{state['analysis_id']}")
        if disclaimer_accepted:
            if state["pdf_bytes"]:
                st.download_button(tr(lang, "download_pdf"), state["pdf_bytes"], file_name=f"{state['analysis_id']}.pdf", mime="application/pdf")
            else:
                st.info("PDF engine is not available in this runtime.")
            if state["docx_bytes"]:
                st.download_button(tr(lang, "download_docx"), state["docx_bytes"], file_name=f"{state['analysis_id']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            json_download(state["package_json"], tr(lang, "download_json"), f"{state['analysis_id']}.json")
            st.download_button(tr(lang, "download_xml"), state["xml_bytes"], file_name=f"{state['analysis_id']}.xml", mime="application/xml")
            st.download_button("Markdown", state["md_report"].encode("utf-8"), file_name=f"{state['analysis_id']}.md", mime="text/markdown")
        else:
            st.info(tr(lang, "downloads_locked"))
        with st.expander(tr(lang, "audit_title"), expanded=False):
            st.json(state["audit"])
            if state["attempt_log"]:
                st.write(pd.DataFrame(state["attempt_log"]))

    with tabs[5]:
        if portfolio_df is None or portfolio_df.empty:
            st.info(tr(lang, "portfolio_placeholder"))
        else:
            pcols = st.columns(4)
            pcols[0].metric(tr(lang, "metrics_assets"), f"{portfolio_summary['count']}")
            pcols[1].metric(tr(lang, "metrics_score"), f"{portfolio_summary['avg_score']}/5")
            pcols[2].metric("High", str(portfolio_summary['high_count']))
            pcols[3].metric("Elevated", str(portfolio_summary['elevated_count']))
            st.caption(f"{market_cfg['currency']} total estimated damage: {format_money(portfolio_summary['total_damage'], market_cfg['currency'])}")
            st.dataframe(portfolio_df, use_container_width=True, hide_index=True)
            dataframe_download(portfolio_df, tr(lang, "download_portfolio"), f"{state['analysis_id']}_portfolio.csv")

else:
    render_section(tr(lang, "results_title"), tr(lang, "results_sub"), tr(lang, "results_kicker"))
    st.info(tr(lang, "need_inputs"))
