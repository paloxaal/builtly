# -*- coding: utf-8 -*-
"""
Builtly | Tender Document Parser
─────────────────────────────────────────────────────────────────
Reell innholdsekstraksjon for anbudsdokumenter. Gir strukturert
output til AI-laget i stedet for bare filnavn-metadata.

Støtter: PDF, DOCX, XLSX/XLS, IFC, DWG/DXF, CSV, TXT, ZIP.
Gjenbrukbar fra andre Builtly-moduler.
"""
from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# ─── Optional parser backends ────────────────────────────────────
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import ifcopenshell
except ImportError:
    ifcopenshell = None

try:
    import ezdxf
except ImportError:
    ezdxf = None


# ─── Classification patterns (two-phase) ─────────────────────────
FILENAME_PATTERNS: Dict[str, List[str]] = {
    "konkurransegrunnlag": [
        r"konkurransegrunnlag", r"kgr\b", r"tender[\s_-]?doc", r"utlysning",
        r"invitation[\s_-]?to[\s_-]?tender", r"itt\b",
    ],
    "beskrivelse": [
        r"beskrivelse", r"kravspek", r"kravspesifikasjon", r"specification",
        r"spesifikasjon", r"teknisk[\s_-]?beskrivelse",
    ],
    "tegning": [
        r"tegning", r"drawing", r"plan[\s_-]?\d", r"snitt", r"fasade[\s_-]?tegning",
        r"arkitekt", r"\bark[\s_-]?\d", r"\brib[\s_-]?\d",
    ],
    "kontrakt": [
        r"kontrakt", r"contract", r"avtale", r"agreement", r"ns[\s_-]?84\d\d",
    ],
    "prisskjema": [
        r"prisskjema", r"pristilbud", r"mengdebeskrivelse", r"mengder",
        r"bid[\s_-]?form", r"price[\s_-]?form", r"bill[\s_-]?of[\s_-]?quantit",
        r"\bboq\b", r"tilbudsskjema",
    ],
    "sha": [
        r"\bsha\b", r"byggherreforskrift", r"sikkerhet[\s_-]?helse",
        r"hms[\s_-]?plan", r"risikovurdering",
    ],
    "miljo": [
        r"milj[øo]plan", r"breeam", r"ceequal", r"ytre[\s_-]?milj",
        r"klimagassregnskap", r"miljoeoppfolging",
    ],
    "rigg": [
        r"rigg[\s_-]?og[\s_-]?drift", r"rigg[\s_-]?plan", r"logistikk",
        r"site[\s_-]?logistics", r"riggomrade",
    ],
    "brann": [
        r"brannkonsept", r"brannstrategi", r"brann[\s_-]?prosjektering",
        r"fire[\s_-]?safety",
    ],
    "geo": [
        r"geoteknisk", r"grunnunders[øo]kelse", r"rig[\s_-]?grunn",
        r"geotechnical",
    ],
    "ifc": [r"\.ifc$"],
    "kvalifikasjon": [
        r"kvalifikasjonskrav", r"qualification", r"esgd", r"\beespd\b",
    ],
    "tildelingskriterier": [
        r"tildelingskriter", r"award[\s_-]?criteria", r"evaluation[\s_-]?criteria",
    ],
}

EXTENSION_HINTS: Dict[str, str] = {
    ".ifc": "ifc",
    ".dwg": "tegning",
    ".dxf": "tegning",
}

# Content signals that override filename classification
CONTENT_SIGNALS: List[Tuple[str, List[str]]] = [
    ("konkurransegrunnlag", [
        "konkurransegrunnlag", "tilbudsfrist", "tildelingskriterier",
        "anskaffelsesform", "kvalifikasjonskrav",
    ]),
    ("kontrakt", [
        "ns 8405", "ns 8406", "ns 8407", "ns8405", "ns8406", "ns8407",
        "dagmulkt", "sikkerhetsstillelse", "kontraktssum", "kontraktsbestemmelser",
    ]),
    ("prisskjema", [
        "mengde", "enhet", "enhetspris", "sum eks mva", "post nr", "rs",
        "kapittelsum", "delsum",
    ]),
    ("sha", [
        "byggherreforskrift", "sha-plan", "sha plan", "risikovurdering",
        "verneombud",
    ]),
    ("brann", [
        "brannkonsept", "brannklasse", "risikoklasse", "rømningsvei",
        "brannmotstand", "rei ", "ei ",
    ]),
    ("geo", [
        "grunnforhold", "borrehull", "kvikkleire", "setning", "fundamenter",
        "geoteknisk vurdering",
    ]),
    ("miljo", [
        "breeam", "klimagass", "ytre miljø", "materialkrav epd",
        "miljøoppfølging",
    ]),
    ("beskrivelse", [
        "kravspesifikasjon", "teknisk beskrivelse", "ytelsesbeskrivelse",
        "funksjonsbeskrivelse",
    ]),
]


def classify_by_filename(name: str) -> str:
    """First-pass classification based on filename only."""
    low = name.lower()
    ext = Path(name).suffix.lower()

    if ext in EXTENSION_HINTS:
        return EXTENSION_HINTS[ext]

    for cat, patterns in FILENAME_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, low):
                return cat
    return "annet"


def classify_by_content(text: str, fallback: str = "annet") -> str:
    """Second-pass classification using content signals."""
    if not text:
        return fallback
    low = text.lower()[:8000]  # first 8k chars is enough

    scores: Dict[str, int] = {}
    for cat, signals in CONTENT_SIGNALS:
        hits = sum(1 for s in signals if s in low)
        if hits > 0:
            scores[cat] = hits

    if not scores:
        return fallback

    best = max(scores.items(), key=lambda kv: kv[1])
    # Only override fallback if content signal is strong
    if best[1] >= 2 or fallback == "annet":
        return best[0]
    return fallback


# ─── Extractors per format ───────────────────────────────────────
def _extract_pdf(data: bytes, max_pages: int = 60) -> Dict[str, Any]:
    """Extract text, tables, and page count from PDF."""
    if not fitz:
        return {"text": "", "page_count": 0, "error": "PyMuPDF ikke installert"}

    out: Dict[str, Any] = {"text": "", "page_count": 0, "tables": []}
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        out["page_count"] = len(doc)
        parts: List[str] = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                parts.append(f"\n\n[... {len(doc) - max_pages} flere sider kuttet ...]\n")
                break
            parts.append(f"\n── Side {i + 1} ──\n{page.get_text('text')}")
        out["text"] = "".join(parts)
        doc.close()
    except Exception as e:
        out["error"] = f"{type(e).__name__}: {e}"

    # Try to extract tables separately with pdfplumber (better for BoQ)
    if pdfplumber and len(data) < 50_000_000:  # 50 MB ceiling
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for pi, page in enumerate(pdf.pages[:max_pages]):
                    tables = page.extract_tables() or []
                    for ti, tbl in enumerate(tables):
                        if tbl and len(tbl) > 1:
                            out["tables"].append({
                                "page": pi + 1,
                                "rows": len(tbl),
                                "cols": len(tbl[0]) if tbl[0] else 0,
                                "preview": tbl[:5],
                            })
        except Exception:
            pass

    return out


def _extract_docx(data: bytes) -> Dict[str, Any]:
    """Extract text and tables from DOCX."""
    if not DocxDocument:
        return {"text": "", "error": "python-docx ikke installert"}

    out: Dict[str, Any] = {"text": "", "tables": []}
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        out["text"] = "\n".join(parts)
        for ti, tbl in enumerate(doc.tables):
            rows = []
            for row in tbl.rows[:30]:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                out["tables"].append({
                    "index": ti,
                    "rows": len(rows),
                    "cols": len(rows[0]) if rows else 0,
                    "preview": rows[:5],
                })
    except Exception as e:
        out["error"] = f"{type(e).__name__}: {e}"
    return out


def _extract_xlsx(data: bytes) -> Dict[str, Any]:
    """Extract sheet structure from XLSX (BoQ, prisskjema etc.)."""
    if not openpyxl:
        return {"text": "", "error": "openpyxl ikke installert"}

    out: Dict[str, Any] = {"text": "", "sheets": []}
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        text_parts: List[str] = []
        for sname in wb.sheetnames[:20]:
            ws = wb[sname]
            sheet_info = {
                "name": sname,
                "max_row": ws.max_row or 0,
                "max_col": ws.max_column or 0,
            }
            # Sample first ~50 rows
            sample_rows: List[List[str]] = []
            for ri, row in enumerate(ws.iter_rows(max_row=50, values_only=True)):
                if ri >= 50:
                    break
                clean_row = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in clean_row):
                    sample_rows.append(clean_row)
            sheet_info["sample"] = sample_rows[:20]
            out["sheets"].append(sheet_info)
            text_parts.append(f"\n── Ark: {sname} ({sheet_info['max_row']}×{sheet_info['max_col']}) ──")
            for row in sample_rows[:30]:
                text_parts.append(" | ".join(row[:10]))
        out["text"] = "\n".join(text_parts)
        wb.close()
    except Exception as e:
        out["error"] = f"{type(e).__name__}: {e}"
    return out


def _extract_ifc(data: bytes) -> Dict[str, Any]:
    """Extract object inventory from IFC."""
    if not ifcopenshell:
        return {"text": "", "error": "ifcopenshell ikke installert"}

    out: Dict[str, Any] = {"text": "", "entity_counts": {}}
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".ifc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        model = ifcopenshell.open(tmp_path)
        # Count entities by type
        counts: Dict[str, int] = {}
        for entity in model:
            t = entity.is_a()
            counts[t] = counts.get(t, 0) + 1
        # Keep top 30 most common
        top = sorted(counts.items(), key=lambda kv: -kv[1])[:30]
        out["entity_counts"] = dict(top)
        lines = [f"IFC schema: {model.schema}"]
        for t, n in top:
            lines.append(f"  {t}: {n}")
        out["text"] = "\n".join(lines)
        Path(tmp_path).unlink(missing_ok=True)
    except Exception as e:
        out["error"] = f"{type(e).__name__}: {e}"
    return out


def _extract_dxf(data: bytes) -> Dict[str, Any]:
    """Extract layer / block summary from DXF (DWG needs conversion)."""
    if not ezdxf:
        return {"text": "", "error": "ezdxf ikke installert"}
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".dxf", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        doc = ezdxf.readfile(tmp_path)
        layers = [layer.dxf.name for layer in doc.layers]
        blocks = [block.name for block in doc.blocks if not block.name.startswith("*")]
        lines = [
            f"DXF version: {doc.dxfversion}",
            f"Antall lag: {len(layers)}",
            f"Antall blokker: {len(blocks)}",
        ]
        if layers:
            lines.append("Lag (topp 30): " + ", ".join(layers[:30]))
        if blocks:
            lines.append("Blokker (topp 30): " + ", ".join(blocks[:30]))
        Path(tmp_path).unlink(missing_ok=True)
        return {"text": "\n".join(lines), "layers": layers[:50], "blocks": blocks[:50]}
    except Exception as e:
        return {"text": "", "error": f"{type(e).__name__}: {e}"}


def _extract_text_file(data: bytes) -> Dict[str, Any]:
    """Plain text / CSV extraction."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return {"text": data.decode(enc)[:100_000]}
        except UnicodeDecodeError:
            continue
    return {"text": "", "error": "Kunne ikke dekode tekstfil"}


def _extract_zip(data: bytes, max_files: int = 20) -> Dict[str, Any]:
    """List contents of a ZIP; do not recurse."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = zf.namelist()[:max_files]
            return {
                "text": "ZIP-innhold:\n" + "\n".join(f"  {n}" for n in names),
                "entries": names,
            }
    except Exception as e:
        return {"text": "", "error": f"ZIP-feil: {e}"}


# ─── Main entry point ────────────────────────────────────────────
def extract_document(filename: str, data: bytes) -> Dict[str, Any]:
    """
    Extract structured content from a single uploaded file.

    Returns:
        {
            "filename": str,
            "extension": str,
            "size_kb": float,
            "category_filename": str,   # first-pass
            "category": str,            # final (content-confirmed)
            "text": str,                # extracted text (possibly truncated)
            "text_excerpt": str,        # first 3000 chars for AI manifest
            "page_count": int,          # PDFs only
            "tables": list,             # structured tables if any
            "sheets": list,             # XLSX sheet info
            "entity_counts": dict,      # IFC
            "error": str | None,
        }
    """
    ext = Path(filename).suffix.lower()
    size_kb = round(len(data) / 1024, 1)
    filename_cat = classify_by_filename(filename)

    extractor_map = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".xlsm": _extract_xlsx,
        ".xls": _extract_xlsx,
        ".ifc": _extract_ifc,
        ".dxf": _extract_dxf,
        ".csv": _extract_text_file,
        ".txt": _extract_text_file,
        ".md": _extract_text_file,
        ".zip": _extract_zip,
    }

    extractor = extractor_map.get(ext)
    if extractor:
        result = extractor(data)
    else:
        result = {"text": "", "error": f"Ingen parser for {ext}"}

    text = result.get("text", "") or ""
    final_cat = classify_by_content(text, fallback=filename_cat)

    return {
        "filename": filename,
        "extension": ext,
        "size_kb": size_kb,
        "category_filename": filename_cat,
        "category": final_cat,
        "text": text,
        "text_excerpt": text[:3000],
        "page_count": result.get("page_count", 0),
        "tables": result.get("tables", []),
        "sheets": result.get("sheets", []),
        "entity_counts": result.get("entity_counts", {}),
        "layers": result.get("layers", []),
        "blocks": result.get("blocks", []),
        "error": result.get("error"),
    }


def extract_documents(files: List[Tuple[str, bytes]]) -> List[Dict[str, Any]]:
    """Batch-extract. `files` is list of (filename, bytes) tuples."""
    return [extract_document(name, data) for name, data in files]


# ─── Quick metadata-extract helpers for specific fields ──────────
DATE_PATTERNS = [
    r"tilbudsfrist[^\n]{0,100}?(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})",
    r"innleveringsfrist[^\n]{0,100}?(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})",
    r"frist for innlevering[^\n]{0,100}?(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})",
    r"submission deadline[^\n]{0,100}?(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})",
]

MONEY_PATTERNS = [
    r"(\d[\d\s\.\,]{3,})\s*(?:kr|nok|mnok|mill(?:ioner)?)",
    r"(?:kr|nok)\s*(\d[\d\s\.\,]{3,})",
]


def quick_scan_deadlines(text: str) -> List[str]:
    """Best-effort deadline extraction — feeds AI but also visible in UI."""
    if not text:
        return []
    found: List[str] = []
    low = text.lower()
    for pat in DATE_PATTERNS:
        for m in re.finditer(pat, low, flags=re.IGNORECASE):
            found.append(m.group(1))
    return list(dict.fromkeys(found))[:10]


def quick_scan_ns_contract(text: str) -> Optional[str]:
    """Detect which NS contract is referenced (8405 / 8406 / 8407)."""
    if not text:
        return None
    low = text.lower()
    for ns in ["ns 8407", "ns8407", "ns 8405", "ns8405", "ns 8406", "ns8406"]:
        if ns in low:
            return ns.upper().replace("NS", "NS ").replace("  ", " ").strip()
    return None


def quick_scan_dagmulkt(text: str) -> Optional[str]:
    """Extract dagmulkt clause snippet."""
    if not text:
        return None
    m = re.search(r"dagmulkt[^\n]{0,200}", text, flags=re.IGNORECASE)
    return m.group(0).strip() if m else None
