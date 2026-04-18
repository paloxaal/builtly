# -*- coding: utf-8 -*-
"""
Builtly Anbudskontroll — Tilbudsbesvarelse
──────────────────────────────────────────────────────────────────────
Genererer utkast til tilbud som skal leveres tilbake til byggherre.
Leverer 7 separate Word-filer som tilbudsansvarlig lett kan bearbeide
og lime inn i egne firmamaler.

Filene:
  01_folgebrev.docx
  02_kvalifikasjonssvar.docx
  03_losningsbeskrivelse.docx
  04_rigg_gjennomforing.docx
  05_sha_hms.docx
  06_fremdriftsplan.docx
  07_forbehold_avvik.docx

Bruker Claude Opus for dypere tekst-generering av høy kvalitet.
"""
from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Cm, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import anthropic
    HAS_CLAUDE = True
except ImportError:
    anthropic = None
    HAS_CLAUDE = False


ANTHROPIC_OPUS_MODEL = os.environ.get("ANTHROPIC_OPUS_MODEL", "claude-opus-4-7")
ANTHROPIC_SONNET_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5")

BUILTLY_NAVY = RGBColor(0x06, 0x11, 0x1A) if HAS_DOCX else None
BUILTLY_BLUE = RGBColor(0x00, 0x60, 0x9B) if HAS_DOCX else None
BUILTLY_GRAY = RGBColor(0x64, 0x74, 0x8B) if HAS_DOCX else None


# ─── AI-promptmaler per seksjon ──────────────────────────────────
SECTION_PROMPTS: Dict[str, Dict[str, str]] = {
    "folgebrev": {
        "title": "Følgebrev",
        "filename": "01_folgebrev.docx",
        "heading": "FØLGEBREV",
        "system": (
            "Du er en tilbudsansvarlig hos en norsk entreprenør som skriver følgebrev til et anbud. "
            "Tonen er profesjonell, konfid, ydmyk overfor byggherren, men selvsikker på egen kompetanse. "
            "Følgebrevet skal: (1) bekrefte innsending og leseforståelse av konkurransegrunnlaget, "
            "(2) kort formidle hvorfor entreprenøren er riktig valg, (3) markere at tilbudet følger med i sin helhet, "
            "(4) inkludere kontaktinfo for oppklaringer. Følgebrevet skal være 300-500 ord. "
            "Skriv som ren prosa i bokmål. Ikke bruk overskrifter eller punktlister. "
            "Ikke bruk floskler som 'vi er stolte av å'."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "OPPDRAGSGIVER: {buyer_name}\n"
            "ENTREPRENØR (avsender): {company_name}\n"
            "KONTRAKTSTYPE: {tender_type}\n"
            "TILBUDSFRIST: {deadline}\n\n"
            "SELSKAPSKONTEKST:\n{company_context}\n\n"
            "KONKURRANSEGRUNNLAG (oppsummert):\n{tender_summary}\n\n"
            "Skriv et følgebrev som {company_name} kan sende med sitt tilbud til {buyer_name}."
        ),
    },

    "kvalifikasjonssvar": {
        "title": "Kvalifikasjonssvar",
        "filename": "02_kvalifikasjonssvar.docx",
        "heading": "SVAR PÅ KVALIFIKASJONSKRAV",
        "system": (
            "Du er en tilbudsansvarlig som besvarer kvalifikasjonskrav i et norsk anbud. "
            "Strukturer svaret med én overskrift per kvalifikasjonskrav fra grunnlaget, og gi "
            "konkrete, verifiserbare svar med referanse til vedlagt dokumentasjon. "
            "Bruk norske fagbegreper. Ikke finn på fakta — hvis informasjon mangler, "
            "marker tydelig '[fyll inn: XYZ]' slik at tilbudsansvarlig kan fullføre manuelt. "
            "Svaret skal struktureres i JSON-array med felter {heading, response_text}."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "OPPDRAGSGIVER: {buyer_name}\n"
            "ENTREPRENØR: {company_name}\n\n"
            "KVALIFIKASJONSKRAV FRA GRUNNLAG:\n{qualification_reqs}\n\n"
            "SELSKAPETS KVALIFIKASJONER:\n"
            "Godkjenningsområder: {approval_areas}\n"
            "Sertifiseringer: {certifications}\n"
            "Referanseprosjekter:\n{reference_projects}\n"
            "HMS-policy: {hms_policy}\n"
            "Kvalitetspolicy: {quality_policy}\n\n"
            "Besvar hvert kvalifikasjonskrav. Output som JSON-array:\n"
            '[{{"heading": "Krav 1 — ...", "response_text": "..."}}]'
        ),
        "json_array": True,
    },

    "losningsbeskrivelse": {
        "title": "Løsningsbeskrivelse",
        "filename": "03_losningsbeskrivelse.docx",
        "heading": "LØSNINGSBESKRIVELSE",
        "system": (
            "Du er tilbudsansvarlig som skriver løsningsbeskrivelse i et norsk bygg/anlegg-anbud. "
            "Beskriv teknisk og prosessuelt hvordan entreprenøren vil løse oppdraget. "
            "Struktur: (1) Overordnet tilnærming, (2) Faglige løsninger per hovedfag, "
            "(3) Kritiske risikoområder og tiltak, (4) Kvalitetssikring og sluttkontroll. "
            "Bruk riktige norske faguttrykk (TEK17, NS-standarder, BIM). "
            "Vær konkret om teknikk, ikke generisk. "
            "Marker [fyll inn: ...] der tilbudsansvarlig må legge inn prosjektspesifikk info. "
            "Skriv som sammenhengende prosa med seksjonsoverskrifter. 1500-2500 ord."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "OPPDRAGSGIVER: {buyer_name}\n"
            "ENTREPRENØR: {company_name}\n"
            "ENTREPRISEFORM: {tender_type}\n\n"
            "PROSJEKTSAMMENDRAG:\n{tender_summary}\n\n"
            "UE-PAKKER (fag):\n{packages}\n\n"
            "VIKTIGE TEKNISKE KRAV:\n{technical_requirements}\n\n"
            "IDENTIFISERTE RISIKOER:\n{risks}\n\n"
            "Skriv løsningsbeskrivelse som bokmål-prosa med seksjoner."
        ),
    },

    "rigg_gjennomforing": {
        "title": "Rigg og gjennomføring",
        "filename": "04_rigg_gjennomforing.docx",
        "heading": "RIGG OG GJENNOMFØRING",
        "system": (
            "Du er prosjekteringsansvarlig som skriver rigg- og gjennomføringsplan i et norsk anbud. "
            "Dekk: (1) Byggeplassorganisasjon og bemanning, (2) Rigg-plan og arbeidsområder, "
            "(3) Logistikk og materiallevering, (4) Koordinering med andre fag/UE, "
            "(5) Fremdriftsstyring og kontraktsmøter, (6) BIM-bruk og digital samhandling. "
            "Bruk norske byggebegreper. Marker [fyll inn: ...] der prosjektspesifikk info trengs. "
            "Skriv som prosa med seksjonsoverskrifter. 1000-1500 ord."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "OPPDRAGSGIVER: {buyer_name}\n"
            "ENTREPRENØR: {company_name}\n"
            "ENTREPRISEFORM: {tender_type}\n\n"
            "PROSJEKTSAMMENDRAG:\n{tender_summary}\n\n"
            "TOMT/SITE-FORHOLD:\n{site_conditions}\n\n"
            "UE-PAKKER:\n{packages}\n\n"
            "Skriv rigg- og gjennomføringsplan som bokmål-prosa."
        ),
    },

    "sha_hms": {
        "title": "SHA og HMS",
        "filename": "05_sha_hms.docx",
        "heading": "SHA OG HMS",
        "system": (
            "Du er SHA-koordinator/HMS-leder som skriver HMS- og SHA-seksjon i et norsk anbud. "
            "Dekk: (1) HMS-policy og -system, (2) SHA-oppfølging i byggefasen, (3) Risikovurdering "
            "og SJA-prosedyre, (4) Opplæring og sertifiseringer, (5) Avviksbehandling og rapportering, "
            "(6) Særskilte SHA-tiltak for dette prosjektet. "
            "Følg byggherreforskriften og relevante NS-standarder. "
            "Marker [fyll inn: ...] der prosjektspesifikt trengs. 800-1200 ord prosa."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "ENTREPRENØR: {company_name}\n"
            "HMS-POLICY: {hms_policy}\n"
            "HMS-SERTIFISERINGER: {certifications}\n\n"
            "SHA-KRAV FRA GRUNNLAG:\n{sha_requirements}\n\n"
            "RISIKOBILDE:\n{risks}\n\n"
            "Skriv SHA/HMS-seksjon som bokmål-prosa."
        ),
    },

    "fremdriftsplan": {
        "title": "Fremdriftsplan",
        "filename": "06_fremdriftsplan.docx",
        "heading": "FREMDRIFTSPLAN",
        "system": (
            "Du er prosjektleder som strukturerer fremdriftsplan i et norsk anbud. "
            "Dette skal være struktur + milepæler (IKKE detaljert Gantt — det gjøres i MS Project separat). "
            "Dekk: (1) Hovedfaser og tidsestimater, (2) Kritiske milepæler, (3) Avhengigheter mellom fag, "
            "(4) Bufferstrategi og risiko-milepæler. "
            "Marker [fyll inn: ...] for datoer som krever innspill fra kalkylestaben. "
            "Struktur: tekst + tabell med milepæler."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "ENTREPRENØR: {company_name}\n\n"
            "VIKTIGE DATOER FRA GRUNNLAG:\n{key_dates}\n\n"
            "UE-PAKKER:\n{packages}\n\n"
            "Skriv fremdriftsplan med milepælstabell som prosa + tabell-beskrivelse.\n"
            "Marker [fyll inn: dato] der du mangler konkret informasjon."
        ),
    },

    "forbehold_avvik": {
        "title": "Forbehold og avvik",
        "filename": "07_forbehold_avvik.docx",
        "heading": "FORBEHOLD OG AVVIK FRA KONKURRANSEGRUNNLAG",
        "system": (
            "Du er tilbudsansvarlig som oppsummerer forbehold og avvik i et norsk anbud. "
            "VIKTIG: Vær konservativ. Ikke finn på forbehold som ikke er identifisert. "
            "Dekk: (1) Avvik fra krav i grunnlaget (hvis noen), (2) Forutsetninger tilbudet bygger på, "
            "(3) Tilbyders kommersielle vilkår (betalingsbetingelser, garantier), "
            "(4) Nøkkelbetingelser ved kontraktsinngåelse. "
            "Hvis ingenting er identifisert, skriv en kort konservativ oversikt som noterer at "
            "tilbudet er uten forbehold og følger konkurransegrunnlaget i sin helhet. "
            "500-1000 ord."
        ),
        "user_template": (
            "PROSJEKT: {project_name}\n"
            "ENTREPRENØR: {company_name}\n\n"
            "IDENTIFISERTE AVVIK/MANGLER FRA PASS2-ANALYSE:\n{pass2_findings}\n\n"
            "KONTRAKTSGRUNNLAG:\n{contract_basis}\n\n"
            "Formuler forbeholds-seksjon som prosa."
        ),
    },
}


# ─── Hoved-AI-kall ───────────────────────────────────────────────
def generate_section_content(
    section_key: str,
    context: Dict[str, str],
    use_opus: bool = True,
) -> Dict[str, Any]:
    """Kjør Claude for én seksjon. Returnerer {text, tokens_in, tokens_out, error}."""
    if not HAS_CLAUDE:
        return {"text": "", "error": "Claude-klient ikke tilgjengelig"}

    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        return {"text": "", "error": "ANTHROPIC_API_KEY ikke satt"}

    prompt = SECTION_PROMPTS.get(section_key)
    if not prompt:
        return {"text": "", "error": f"Ukjent seksjon: {section_key}"}

    # Fyll inn template
    try:
        user_msg = prompt["user_template"].format(**context)
    except KeyError as e:
        return {"text": "", "error": f"Mangler kontekst-felt: {e}"}

    model = ANTHROPIC_OPUS_MODEL if use_opus else ANTHROPIC_SONNET_MODEL

    try:
        client = anthropic.Anthropic(api_key=key)
        resp = client.messages.create(
            model=model,
            max_tokens=4000,
            temperature=0.3,
            system=prompt["system"],
            messages=[{"role": "user", "content": user_msg}],
        )
        text = resp.content[0].text if resp.content else ""
        return {
            "text": text,
            "is_json_array": prompt.get("json_array", False),
            "tokens_in": getattr(resp.usage, "input_tokens", 0),
            "tokens_out": getattr(resp.usage, "output_tokens", 0),
            "model": model,
        }
    except Exception as e:
        return {"text": "", "error": f"{type(e).__name__}: {str(e)[:200]}"}


# ─── DOCX-generator ──────────────────────────────────────────────
def generate_section_docx(
    section_key: str,
    section_text: str,
    company_profile: Dict[str, Any],
    project_name: str,
    buyer_name: str = "",
    is_json_array: bool = False,
) -> bytes:
    """Bygg en enkelt seksjons-DOCX med firmabrev-header."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx ikke tilgjengelig")

    prompt = SECTION_PROMPTS.get(section_key, {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Firmabrev-header
    if company_profile.get("company_name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(company_profile["company_name"])
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = BUILTLY_NAVY

        address_parts = []
        if company_profile.get("company_address"):
            address_parts.append(company_profile["company_address"])
        pc_city = " ".join(x for x in [
            company_profile.get("company_postcode", ""),
            company_profile.get("company_city", ""),
        ] if x).strip()
        if pc_city:
            address_parts.append(pc_city)
        if company_profile.get("company_org_no"):
            address_parts.append(f"Org.nr.: {company_profile['company_org_no']}")

        if address_parts:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            addr_run = p2.add_run(" · ".join(address_parts))
            addr_run.font.size = Pt(9)
            addr_run.font.color.rgb = BUILTLY_GRAY

    # Tittel
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(prompt.get("heading", section_key.upper()))
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BUILTLY_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"{project_name}" + (f" — {buyer_name}" if buyer_name else ""))
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = BUILTLY_BLUE

    doc.add_paragraph()

    # Innhold — håndter JSON-array (f.eks. kvalifikasjonssvar) annerledes enn ren prosa
    if is_json_array:
        _render_json_array_content(doc, section_text)
    else:
        _render_prose_content(doc, section_text)

    # Signatur
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(f"{datetime.now().strftime('%d.%m.%Y')}").italic = True
    sig.add_run("\n\n")
    sig.add_run("Med vennlig hilsen\n\n").italic = True

    contact_name = company_profile.get("contact_person", "")
    if contact_name:
        name_run = sig.add_run(f"{contact_name}\n")
        name_run.bold = True
    if company_profile.get("contact_title"):
        sig.add_run(f"{company_profile['contact_title']}\n")
    if company_profile.get("company_name"):
        sig.add_run(company_profile["company_name"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_prose_content(doc, text: str):
    """Rendre ren prosa-tekst, tolker markdown-lignende headers (## og **bold**)."""
    if not text:
        doc.add_paragraph("[Innhold ikke generert]")
        return

    lines = text.split("\n")
    current_para = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_para is not None:
                current_para = None
            continue

        # Markdown-header?
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = BUILTLY_BLUE
            current_para = None
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:])
            r.bold = True
            r.font.size = Pt(12.5)
            r.font.color.rgb = BUILTLY_NAVY
            current_para = None
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = BUILTLY_NAVY
            current_para = None
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            current_para = None
        elif re.match(r"^\d+[.)]\s", stripped):
            # Nummerert liste
            p = doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
            current_para = None
        else:
            # Ren tekst — slå sammen sammenhengende linjer til ett avsnitt
            if current_para is None:
                current_para = doc.add_paragraph()
            else:
                current_para.add_run(" ")
            # Håndter **bold** i teksten
            _add_text_with_markdown_bold(current_para, stripped)


def _add_text_with_markdown_bold(para, text: str):
    """Parse **bold** og legg til runs med riktig formatering."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def _render_json_array_content(doc, text: str):
    """Tolk AI-respons som JSON-array med {heading, response_text}-objekter."""
    # Plukk ut JSON-array
    t = text.strip()
    t = re.sub(r"^```(?:json)?", "", t, flags=re.IGNORECASE).strip()
    t = re.sub(r"```$", "", t).strip()

    first = t.find("[")
    last = t.rfind("]")
    if first == -1 or last == -1:
        # Fallback — render som prosa
        _render_prose_content(doc, text)
        return

    try:
        items = json.loads(t[first:last + 1])
    except Exception:
        _render_prose_content(doc, text)
        return

    if not isinstance(items, list):
        _render_prose_content(doc, text)
        return

    for item in items:
        if not isinstance(item, dict):
            continue
        heading = item.get("heading", "")
        response = item.get("response_text", "")

        if heading:
            p = doc.add_paragraph()
            r = p.add_run(heading)
            r.bold = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = BUILTLY_NAVY

        if response:
            for para_text in response.split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    _add_text_with_markdown_bold(p, para_text.strip())


# ─── Orchestrator: generer alle 7 seksjoner ──────────────────────
def generate_all_response_sections(
    company_profile: Dict[str, Any],
    project_name: str,
    buyer_name: str,
    tender_type: str,
    pass1_data: List[Dict[str, Any]],
    pass2_data: Dict[str, Any],
    packages: List[str],
    deadline: Optional[str] = None,
    use_opus: bool = True,
    progress_callback=None,
) -> List[Dict[str, Any]]:
    """
    Generer alle 7 DOCX-seksjoner for tilbudsbesvarelse.

    Returnerer liste av:
      { section_key, filename, bytes, error, tokens_in, tokens_out }
    """
    # Bygg delt kontekst fra pass-data
    shared_context = _build_shared_context(
        company_profile=company_profile,
        project_name=project_name,
        buyer_name=buyer_name,
        tender_type=tender_type,
        pass1_data=pass1_data,
        pass2_data=pass2_data,
        packages=packages,
        deadline=deadline,
    )

    results = []
    sections = list(SECTION_PROMPTS.keys())

    for i, section_key in enumerate(sections, 1):
        if progress_callback:
            try:
                progress_callback(i, len(sections), section_key)
            except Exception:
                pass

        prompt = SECTION_PROMPTS[section_key]
        content_result = generate_section_content(
            section_key=section_key,
            context=shared_context,
            use_opus=use_opus,
        )

        if content_result.get("error"):
            results.append({
                "section_key": section_key,
                "title": prompt["title"],
                "filename": prompt["filename"],
                "bytes": None,
                "error": content_result["error"],
                "tokens_in": 0,
                "tokens_out": 0,
            })
            continue

        try:
            docx_bytes = generate_section_docx(
                section_key=section_key,
                section_text=content_result["text"],
                company_profile=company_profile,
                project_name=project_name,
                buyer_name=buyer_name,
                is_json_array=content_result.get("is_json_array", False),
            )
            results.append({
                "section_key": section_key,
                "title": prompt["title"],
                "filename": prompt["filename"],
                "bytes": docx_bytes,
                "error": None,
                "tokens_in": content_result.get("tokens_in", 0),
                "tokens_out": content_result.get("tokens_out", 0),
            })
        except Exception as e:
            results.append({
                "section_key": section_key,
                "title": prompt["title"],
                "filename": prompt["filename"],
                "bytes": None,
                "error": f"DOCX-generering feilet: {type(e).__name__}: {e}",
                "tokens_in": content_result.get("tokens_in", 0),
                "tokens_out": content_result.get("tokens_out", 0),
            })

    return results


def _build_shared_context(
    company_profile: Dict[str, Any],
    project_name: str,
    buyer_name: str,
    tender_type: str,
    pass1_data: List[Dict[str, Any]],
    pass2_data: Dict[str, Any],
    packages: List[str],
    deadline: Optional[str],
) -> Dict[str, str]:
    """Bygg strengifisert kontekst for promptfelt."""
    # Pass1 — utdrag fra dokumenter
    doc_summary = []
    for d in pass1_data[:20]:
        excerpt = (d.get("text_excerpt") or d.get("text", ""))[:1500]
        if excerpt:
            doc_summary.append(f"[{d.get('filename', '?')} ({d.get('category', '?')})]\n{excerpt}")
    tender_summary = "\n\n".join(doc_summary)[:15000]

    # Referanseprosjekter
    ref_lines = []
    for ref in (company_profile.get("reference_projects") or [])[:10]:
        if isinstance(ref, dict):
            parts = [ref.get("name", "")]
            if ref.get("value_mnok"):
                parts.append(f"{ref['value_mnok']} MNOK")
            if ref.get("year"):
                parts.append(str(ref["year"]))
            if ref.get("role"):
                parts.append(ref["role"])
            ref_lines.append("  - " + " · ".join(str(p) for p in parts if p))
            if ref.get("description"):
                ref_lines.append(f"    {ref['description']}")
        elif isinstance(ref, str):
            ref_lines.append(f"  - {ref}")
    reference_projects = "\n".join(ref_lines) or "[Ingen referanseprosjekter i profil]"

    # Pass2-felt
    qualification_reqs = pass2_data.get("qualification_requirements_raw", "[Ikke identifisert]") or "[Ikke identifisert]"
    if isinstance(qualification_reqs, list):
        qualification_reqs = "\n".join(f"- {q}" for q in qualification_reqs)
    sha_requirements = pass2_data.get("sha_requirements", "[Ikke spesifisert i grunnlaget]") or "[Ikke spesifisert i grunnlaget]"
    technical_requirements = pass2_data.get("technical_requirements_summary", "[Se konkurransegrunnlag]") or "[Se konkurransegrunnlag]"
    if isinstance(technical_requirements, list):
        technical_requirements = "\n".join(f"- {t}" for t in technical_requirements)

    risks_list = pass2_data.get("risks") or []
    risks = "\n".join(
        f"- [{r.get('severity', '?')}] {r.get('description', '')}"
        for r in risks_list[:15]
        if isinstance(r, dict)
    ) or "[Ingen eksplisitte risikoer fra pass2]"

    key_dates = pass2_data.get("key_dates") or []
    key_dates_str = "\n".join(
        f"- {d.get('label', '')}: {d.get('date', '')}"
        for d in key_dates[:15]
        if isinstance(d, dict)
    ) or "[Se grunnlag]"

    pass2_findings = pass2_data.get("findings_summary") or pass2_data.get("cross_check_summary") or "[Ingen funn]"
    if isinstance(pass2_findings, list):
        pass2_findings = "\n".join(f"- {f}" for f in pass2_findings[:20])

    contract_basis = pass2_data.get("contract_basis", "NS 8405/8406/8407 — se grunnlag")
    site_conditions = pass2_data.get("site_conditions", "[Se konkurransegrunnlag]")

    # Selskaps-kontekst
    company_context = []
    if company_profile.get("company_description"):
        company_context.append(company_profile["company_description"])
    if company_profile.get("approval_areas"):
        company_context.append(
            "Godkjenningsområder: " + ", ".join(company_profile["approval_areas"])
        )
    if company_profile.get("certifications"):
        company_context.append(
            "Sertifiseringer: " + ", ".join(company_profile["certifications"])
        )
    company_context_str = "\n".join(company_context) or "[Ingen selskapsbeskrivelse]"

    return {
        "company_name": company_profile.get("company_name", ""),
        "project_name": project_name,
        "buyer_name": buyer_name,
        "tender_type": tender_type,
        "deadline": deadline or "[ikke spesifisert]",
        "company_context": company_context_str,
        "tender_summary": tender_summary,
        "qualification_reqs": qualification_reqs,
        "approval_areas": ", ".join(company_profile.get("approval_areas") or []) or "[ingen]",
        "certifications": ", ".join(company_profile.get("certifications") or []) or "[ingen]",
        "reference_projects": reference_projects,
        "hms_policy": company_profile.get("hms_policy", "[Ikke satt i profil]"),
        "quality_policy": company_profile.get("quality_policy", "[Ikke satt i profil]"),
        "packages": "\n".join(f"- {p}" for p in packages) or "[ikke spesifisert]",
        "technical_requirements": technical_requirements,
        "risks": risks,
        "site_conditions": site_conditions,
        "sha_requirements": sha_requirements,
        "key_dates": key_dates_str,
        "pass2_findings": pass2_findings,
        "contract_basis": contract_basis,
    }


# ─── ZIP-bygger for nedlasting ───────────────────────────────────
def build_response_zip(results: List[Dict[str, Any]], project_name: str = "") -> bytes:
    """Pakk alle genererte DOCX-filer i én ZIP."""
    import zipfile
    safe_project = re.sub(r"[^\w\-]", "_", project_name)[:40] or "tilbud"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for r in results:
            if r.get("bytes") and r.get("filename"):
                zf.writestr(r["filename"], r["bytes"])

        # Legg ved en README
        readme = f"""TILBUDSBESVARELSE — {project_name}
Generert av Builtly Anbudskontroll: {datetime.now().strftime('%d.%m.%Y %H:%M')}

Filer i denne mappen:
""" + "\n".join(f"  {r['filename']} — {r['title']}" for r in results if r.get("bytes"))
        readme += """

Merk: Dette er AI-genererte utkast som skal gjennomgås og tilpasses av tilbudsansvarlig
før levering. Alle sted hvor det står [fyll inn: ...] krever manuell utfylling.
"""
        zf.writestr("README.txt", readme.encode("utf-8"))

    return buf.getvalue()


def response_module_selfcheck() -> Dict[str, Any]:
    return {
        "has_docx": HAS_DOCX,
        "has_claude": HAS_CLAUDE,
        "claude_api_key_set": bool(os.environ.get("ANTHROPIC_API_KEY")),
        "opus_model": ANTHROPIC_OPUS_MODEL,
        "sonnet_model": ANTHROPIC_SONNET_MODEL,
    }
