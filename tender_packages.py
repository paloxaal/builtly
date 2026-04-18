# -*- coding: utf-8 -*-
"""
Builtly Anbudskontroll — UE-tilbudsgrunnlag
──────────────────────────────────────────────────────────────────────
Genererer komplett tilbudsgrunnlag til underentreprenører per fag/pakke.

Flyt:
  1. Henter inn pass1/pass2-data fra eksisterende anbudsanalyse
  2. AI trekker ut fagspesifikt innhold per UE-pakke
  3. Genererer DOCX via python-docx med 9 seksjoner
  4. Returnerer både filsti og bytes for nedlasting

Input forventer at full_analysis-strukturen finnes i session_state:
  - pass1: per-dokument analyse
  - pass2: krysskontroll, mangler, forutsetninger
  - packages: liste av UE-pakker fra intake
"""
from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Cm, RGBColor, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import anthropic
    HAS_CLAUDE = True
except ImportError:
    anthropic = None
    HAS_CLAUDE = False


ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5")


# ─── AI-kall for per-pakke-ekstraksjon ──────────────────────────
PACKAGE_EXTRACTION_SYSTEM = """Du er en tilbudsansvarlig hos en norsk totalentreprenør som skal
lage tilbudsgrunnlag til underentreprenører. Jobben din er å trekke ut alt som er relevant for
ÉN spesifikk UE-pakke fra et fullstendig konkurransegrunnlag.

Vær konkret og presis. Bruk norske faguttrykk (RIB, RIV, RIE, ARK, TEK17, SHA osv.).
Ikke oppfinn tall eller krav som ikke står i grunnlaget — hvis noe mangler,
merk det som "[ikke spesifisert i grunnlaget]".

Output skal være JSON med denne strukturen:
{
  "scope_description": "2-4 avsnitt som beskriver hva UE skal levere. Konkret og målrettet.",
  "technical_requirements": [
    {"reference": "TEK17 § X-Y / NS X", "description": "konkret krav"}
  ],
  "quantities": [
    {"post": "post-nr hvis finnes", "description": "kort beskrivelse", "unit": "m², lm, stk, rs", "quantity": "tall eller tekst"}
  ],
  "interfaces": [
    {"against_discipline": "f.eks. RIB eller ARK", "description": "hva UE må koordinere"}
  ],
  "site_conditions": "Rigg, logistikk, tilkomst, fellesytelser — kort",
  "hms_requirements": "SHA-krav som gjelder denne UE-en spesielt",
  "special_conditions": "Forbehold, særskilte krav, spesialmaterialer"
}

Svar KUN med JSON-objektet. Ingen forklarende tekst før eller etter."""


def extract_package_content(
    package_name: str,
    packages_all: List[str],
    pass1_data: List[Dict[str, Any]],
    pass2_data: Dict[str, Any],
    tender_type: str = "Totalentreprise",
    project_name: str = "",
) -> Dict[str, Any]:
    """Kjør AI-ekstraksjon av UE-relevant innhold for én pakke."""
    if not HAS_CLAUDE:
        return _empty_package_content(f"Claude-klient ikke tilgjengelig")

    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        return _empty_package_content("ANTHROPIC_API_KEY ikke satt")

    # Bygg konsentrert input fra pass1-data (tekstutdrag per dokument)
    doc_excerpts = []
    for d in pass1_data[:30]:  # cap for å unngå context-eksplosjon
        excerpt = d.get("text_excerpt") or (d.get("text", "")[:2000])
        if excerpt:
            doc_excerpts.append(
                f"--- {d.get('filename', 'ukjent')} ({d.get('category', 'annet')}) ---\n{excerpt}"
            )

    # Bygg kontekst fra pass2-data
    pass2_context = ""
    if pass2_data.get("key_dates"):
        pass2_context += "Viktige datoer:\n"
        for dt in pass2_data["key_dates"][:10]:
            pass2_context += f"  - {dt.get('label', '')}: {dt.get('date', '')}\n"
    if pass2_data.get("contract_basis"):
        pass2_context += f"\nKontraktsgrunnlag: {pass2_data['contract_basis']}\n"

    user_msg = (
        f"PROSJEKT: {project_name}\n"
        f"ANSKAFFELSESFORM: {tender_type}\n\n"
        f"UE-PAKKE Å EKSTRAHERE INNHOLD FOR: {package_name}\n\n"
        f"ALLE PAKKER I PROSJEKTET: {', '.join(packages_all)}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        f"KONKURRANSEGRUNNLAG (utdrag per dokument):\n\n"
        + "\n\n".join(doc_excerpts)
        + f"\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        + f"KRYSSANALYSE:\n{pass2_context}\n"
        + f"\nTrekk ut alt som gjelder UE-pakken '{package_name}' spesielt.\n"
        + f"Vær presis med tekniske krav, grensesnitt mot andre fag, og mengder der de finnes."
    )

    try:
        client = anthropic.Anthropic(api_key=key)
        resp = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=3000,
            temperature=0.15,
            system=PACKAGE_EXTRACTION_SYSTEM,
            messages=[{"role": "user", "content": user_msg}],
        )
        text = resp.content[0].text if resp.content else ""
        parsed = _parse_json(text)
        if not parsed:
            return _empty_package_content("Kunne ikke parse AI-respons")

        # Normaliser felt
        return {
            "scope_description": parsed.get("scope_description", ""),
            "technical_requirements": parsed.get("technical_requirements", []) or [],
            "quantities": parsed.get("quantities", []) or [],
            "interfaces": parsed.get("interfaces", []) or [],
            "site_conditions": parsed.get("site_conditions", ""),
            "hms_requirements": parsed.get("hms_requirements", ""),
            "special_conditions": parsed.get("special_conditions", ""),
            "tokens_in": getattr(resp.usage, "input_tokens", 0),
            "tokens_out": getattr(resp.usage, "output_tokens", 0),
        }
    except Exception as e:
        return _empty_package_content(f"AI-feil: {type(e).__name__}: {str(e)[:100]}")


def _empty_package_content(error_note: str = "") -> Dict[str, Any]:
    return {
        "scope_description": f"[Ekstraksjon feilet: {error_note}]" if error_note else "",
        "technical_requirements": [],
        "quantities": [],
        "interfaces": [],
        "site_conditions": "",
        "hms_requirements": "",
        "special_conditions": "",
        "error": error_note,
    }


def _parse_json(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    t = text.strip()
    t = re.sub(r"^```(?:json)?", "", t, flags=re.IGNORECASE).strip()
    t = re.sub(r"```$", "", t).strip()
    first = t.find("{")
    last = t.rfind("}")
    if first == -1 or last == -1:
        return None
    try:
        return json.loads(t[first:last + 1])
    except Exception:
        return None


# ─── DOCX-generator ──────────────────────────────────────────────
BUILTLY_NAVY = RGBColor(0x06, 0x11, 0x1A) if HAS_DOCX else None
BUILTLY_BLUE = RGBColor(0x00, 0x60, 0x9B) if HAS_DOCX else None
BUILTLY_GRAY = RGBColor(0x64, 0x74, 0x8B) if HAS_DOCX else None


def _set_cell_background(cell, color_hex: str):
    """Hjelpefunksjon for å sette bakgrunnsfarge på tabellcelle."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def generate_package_docx(
    package_name: str,
    package_content: Dict[str, Any],
    company_profile: Dict[str, Any],
    project_name: str,
    tender_type: str,
    deadline: Optional[str] = None,
    response_deadline: Optional[str] = None,
    relevant_drawings: Optional[List[Dict[str, Any]]] = None,
    buyer_name: str = "",
) -> bytes:
    """
    Generer en full DOCX for en UE-tilbudsforespørsel med 9 seksjoner.

    Returnerer DOCX-innhold som bytes.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx ikke tilgjengelig")

    doc = Document()

    # Sett standard font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Sett sidemarginer
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── FIRMABREV-HEADER ────────────────────────────────────────
    _add_letterhead(doc, company_profile)

    # ─── DOKUMENTTITTEL ──────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TILBUDSFORESPØRSEL — UNDERENTREPRISE")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = BUILTLY_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Pakke: {package_name}")
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = BUILTLY_BLUE

    # ─── METADATA-TABELL ─────────────────────────────────────────
    doc.add_paragraph()
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Prosjekt", project_name or "[navn]"),
        ("Oppdragsgiver (byggherre)", buyer_name or "[navn]"),
        ("Totalentreprenør", company_profile.get("company_name", "")),
        ("Entrepriseform (hovedkontrakt)", tender_type),
        ("Dato", datetime.now().strftime("%d.%m.%Y")),
    ]
    for i, (k, v) in enumerate(meta_rows):
        row = meta_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_background(row.cells[0], "F1F5F9")

    doc.add_paragraph()

    # ─── 1. BAKGRUNN OG PROSJEKTINFO ─────────────────────────────
    _add_heading(doc, "1. Bakgrunn og prosjektinformasjon")
    p = doc.add_paragraph()
    p.add_run(
        f"{company_profile.get('company_name', '[totalentreprenør]')} er totalentreprenør "
        f"for prosjekt {project_name or '[prosjektnavn]'}"
        + (f" for {buyer_name}" if buyer_name else "")
        + f". Vi inviterer med dette til å avgi tilbud på UE-pakken "
        f"«{package_name}»."
    )

    if company_profile.get("company_description"):
        doc.add_paragraph(company_profile["company_description"])

    # ─── 2. LEVERANSEOMFANG ──────────────────────────────────────
    _add_heading(doc, "2. Leveranseomfang")
    scope = package_content.get("scope_description") or f"[Scope for {package_name} — fyll ut manuelt]"
    for para in scope.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ─── 3. TEKNISKE KRAV OG REFERANSER ──────────────────────────
    _add_heading(doc, "3. Tekniske krav og referanser")
    tech_reqs = package_content.get("technical_requirements", [])
    if tech_reqs:
        tech_table = doc.add_table(rows=1, cols=2)
        tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tech_table.rows[0].cells
        hdr[0].text = "Referanse"
        hdr[1].text = "Krav"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_background(cell, "06111A")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for req in tech_reqs:
            row = tech_table.add_row().cells
            row[0].text = req.get("reference", "")
            row[1].text = req.get("description", "")
    else:
        doc.add_paragraph(
            "[Tekniske krav skal hentes fra konkurransegrunnlag — "
            "komplett referanseliste følger som vedlegg.]"
        )

    # ─── 4. MENGDER ──────────────────────────────────────────────
    _add_heading(doc, "4. Mengder")
    quantities = package_content.get("quantities", [])
    if quantities:
        qty_table = doc.add_table(rows=1, cols=4)
        qty_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = qty_table.rows[0].cells
        hdr[0].text = "Post"
        hdr[1].text = "Beskrivelse"
        hdr[2].text = "Enhet"
        hdr[3].text = "Mengde"
        for cell in hdr:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_background(cell, "06111A")
        for q in quantities:
            row = qty_table.add_row().cells
            row[0].text = str(q.get("post", ""))
            row[1].text = str(q.get("description", ""))
            row[2].text = str(q.get("unit", ""))
            row[3].text = str(q.get("quantity", ""))
        doc.add_paragraph()
        note = doc.add_paragraph()
        note_run = note.add_run(
            "Merk: Fullstendig mengdefortegnelse finnes i vedlagt prisskjema. "
            "UE bekrefter mengder ved befaring og ved gjennomgang av tegninger."
        )
        note_run.italic = True
        note_run.font.size = Pt(9.5)
        note_run.font.color.rgb = BUILTLY_GRAY
    else:
        doc.add_paragraph(
            "Mengdefortegnelse: Se vedlagt prisskjema. UE er ansvarlig for å kontrollregne mengder."
        )

    # ─── 5. GRENSESNITT MOT ANDRE FAG ────────────────────────────
    _add_heading(doc, "5. Grensesnitt mot andre fag")
    interfaces = package_content.get("interfaces", [])
    if interfaces:
        for iface in interfaces:
            p = doc.add_paragraph(style="List Bullet")
            run1 = p.add_run(f"Mot {iface.get('against_discipline', 'annet fag')}: ")
            run1.bold = True
            p.add_run(iface.get("description", ""))
    else:
        doc.add_paragraph(
            "Grensesnitt mellom fag skal koordineres gjennom prosjektets "
            "koordineringsmøter. UE forplikter seg til å delta på fagmøter "
            "etter innkalling."
        )

    # ─── 6. RIGG, HMS, SHA-KRAV ──────────────────────────────────
    _add_heading(doc, "6. Rigg, HMS og SHA-krav")
    site = package_content.get("site_conditions")
    if site:
        doc.add_paragraph("Rigg og logistikk:").runs[0].bold = True
        doc.add_paragraph(site)
    hms = package_content.get("hms_requirements")
    if hms:
        doc.add_paragraph("HMS- og SHA-krav:").runs[0].bold = True
        doc.add_paragraph(hms)
    doc.add_paragraph(
        "UE skal dokumentere HMS-kompetanse, sende inn SJA for kritiske operasjoner, "
        "og følge byggherres SHA-plan. Tilbyder bekrefter i tilbudet at man har et "
        "etablert HMS-system og skal rapportere som avtalt i hovedkontrakten."
    )

    # ─── 7. TILBUDSFORMAT OG FRISTER ─────────────────────────────
    _add_heading(doc, "7. Tilbudsformat og frister")
    format_lines = [
        "Tilbudet skal inneholde:",
        "  • Prisskjema utfylt med enhetspris og sum per post",
        "  • Bekreftet leveranseomfang i henhold til pkt. 2",
        "  • Forbehold og tilbudsforutsetninger eksplisitt listet",
        "  • Fremdriftsplan og bemanning",
        "  • Nøkkelpersonell med CV (prosjektleder, anleggsleder)",
        "  • Referanser fra lignende prosjekter siste 3 år",
        "  • HMS-dokumentasjon (rammeverk, SJA-eksempel)",
        "  • Firmaattest og skatteattest",
    ]
    for line in format_lines:
        doc.add_paragraph(line)

    if response_deadline:
        doc.add_paragraph()
        dl = doc.add_paragraph()
        dl_run = dl.add_run(f"Tilbudsfrist: {response_deadline}")
        dl_run.bold = True
        dl_run.font.size = Pt(11.5)
        dl_run.font.color.rgb = BUILTLY_BLUE

    if deadline:
        doc.add_paragraph(f"Ønsket kontraktsinngåelse / oppstart: {deadline}")

    # ─── 8. RELEVANTE TEGNINGER (manuell avkrysning) ─────────────
    _add_heading(doc, "8. Relevante tegninger og dokumenter")
    if relevant_drawings:
        doc.add_paragraph("Følgende dokumenter følger vedlagt tilbudsforespørselen:")
        for d in relevant_drawings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(d.get("filename", ""))
            if d.get("discipline"):
                p.add_run(f"  ({d['discipline']})").italic = True
    else:
        doc.add_paragraph(
            "[Relevante tegninger og vedlegg: Fyll inn liste over tegninger som sendes "
            "sammen med tilbudsforespørselen — typisk plan/snitt/detaljer for gjeldende fag, "
            "IFC-modell eller ekstrakter, og relevante beskrivelsestekster.]"
        )

    # ─── 9. KONTAKT OG SIGNATUR ──────────────────────────────────
    _add_heading(doc, "9. Kontakt og videre prosess")
    contact_name = company_profile.get("contact_person", "[kontaktperson]")
    contact_email = company_profile.get("contact_email", "[e-post]")
    contact_phone = company_profile.get("contact_phone", "[telefon]")
    doc.add_paragraph(
        f"Spørsmål til tilbudsforespørselen rettes til:"
    )
    p = doc.add_paragraph()
    p.add_run(f"{contact_name}\n").bold = True
    if company_profile.get("contact_title"):
        p.add_run(f"{company_profile['contact_title']}\n")
    p.add_run(f"E-post: {contact_email}\n")
    p.add_run(f"Telefon: {contact_phone}\n")

    doc.add_paragraph(
        "Befaring kan avtales på forespørsel. Tilbudet sendes elektronisk til kontaktperson over. "
        "Tilbud som mottas etter frist forbeholdes avvist."
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Med vennlig hilsen\n\n").italic = True
    sig_name = sig.add_run(f"{contact_name}\n")
    sig_name.bold = True
    if company_profile.get("contact_title"):
        sig.add_run(f"{company_profile['contact_title']}\n")
    sig.add_run(company_profile.get("company_name", ""))

    # ─── Special conditions som sluttnotis ───────────────────────
    if package_content.get("special_conditions"):
        doc.add_paragraph()
        note_heading = doc.add_paragraph()
        note_heading.add_run("Særskilte forhold:").bold = True
        doc.add_paragraph(package_content["special_conditions"])

    # ─── Serialiser til bytes ────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_letterhead(doc, profile: Dict[str, Any]):
    """Legg til firmabrev-header øverst i dokumentet."""
    if not profile.get("company_name"):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(profile["company_name"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BUILTLY_NAVY

    # Adresse og kontakt
    address_parts = []
    if profile.get("company_address"):
        address_parts.append(profile["company_address"])
    pc_city = " ".join(x for x in [
        profile.get("company_postcode", ""),
        profile.get("company_city", ""),
    ] if x).strip()
    if pc_city:
        address_parts.append(pc_city)
    if profile.get("company_org_no"):
        address_parts.append(f"Org.nr.: {profile['company_org_no']}")

    if address_parts:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        addr_run = p2.add_run(" · ".join(address_parts))
        addr_run.font.size = Pt(9)
        addr_run.font.color.rgb = BUILTLY_GRAY

    # Skillelinje
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = p3.add_run("─" * 70)
    line_run.font.color.rgb = BUILTLY_GRAY
    line_run.font.size = Pt(8)


def _add_heading(doc, text: str):
    """Legg til seksjonsoverskrift med Builtly-farger."""
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = BUILTLY_NAVY


# ─── Orchestrator: generer alle pakker på én gang ────────────────
def generate_all_packages(
    packages: List[str],
    pass1_data: List[Dict[str, Any]],
    pass2_data: Dict[str, Any],
    company_profile: Dict[str, Any],
    project_name: str,
    tender_type: str,
    buyer_name: str = "",
    deadline: Optional[str] = None,
    response_deadline: Optional[str] = None,
    drawings_by_package: Optional[Dict[str, List[Dict[str, Any]]]] = None,
    progress_callback=None,
) -> List[Dict[str, Any]]:
    """
    Generer DOCX for alle UE-pakker.

    Returnerer liste av dict med:
      { package_name, filename, bytes, extraction_metadata }
    """
    drawings_by_package = drawings_by_package or {}
    results = []

    for i, package in enumerate(packages, 1):
        if progress_callback:
            try:
                progress_callback(i, len(packages), package)
            except Exception:
                pass

        # 1. AI-ekstraksjon
        content = extract_package_content(
            package_name=package,
            packages_all=packages,
            pass1_data=pass1_data,
            pass2_data=pass2_data,
            tender_type=tender_type,
            project_name=project_name,
        )

        # 2. DOCX-generering
        try:
            docx_bytes = generate_package_docx(
                package_name=package,
                package_content=content,
                company_profile=company_profile,
                project_name=project_name,
                tender_type=tender_type,
                deadline=deadline,
                response_deadline=response_deadline,
                relevant_drawings=drawings_by_package.get(package, []),
                buyer_name=buyer_name,
            )
            safe_package = re.sub(r"[^\w\-]", "_", package)[:50]
            filename = f"UE-tilbudsgrunnlag_{safe_package}.docx"

            results.append({
                "package_name": package,
                "filename": filename,
                "bytes": docx_bytes,
                "extraction_metadata": {
                    "tokens_in": content.get("tokens_in", 0),
                    "tokens_out": content.get("tokens_out", 0),
                    "error": content.get("error"),
                },
            })
        except Exception as e:
            results.append({
                "package_name": package,
                "filename": None,
                "bytes": None,
                "error": f"DOCX-generering feilet: {type(e).__name__}: {e}",
            })

    return results


def estimate_cost_per_package() -> float:
    """Estimert AI-kost per pakke i NOK."""
    # Sonnet 4.5: ~$3/M input, $15/M output. Per pakke: ~4k input, 1k output
    # Kost per pakke: 4000/1M * $3 + 1000/1M * $15 = $0.012 + $0.015 = $0.027 = ~0.3 NOK
    # La oss være konservative og si 0.5 NOK per pakke
    return 0.5


def package_module_selfcheck() -> Dict[str, Any]:
    """Returner status for å sjekke at modulen er klar."""
    return {
        "has_docx": HAS_DOCX,
        "has_claude": HAS_CLAUDE,
        "claude_api_key_set": bool(os.environ.get("ANTHROPIC_API_KEY")),
        "model": ANTHROPIC_MODEL,
    }
