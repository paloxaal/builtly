# -*- coding: utf-8 -*-
"""
Builtly | Tender Pricing Packages
─────────────────────────────────────────────────────────────────
Genererer underlagspakker per fag/pakke for utsendelse til
underentreprenører for ekstern prising. Produserer DOCX-forespørsel
med scope, krav, frister og svarfelter, og integrerer mot Resend
for faktisk utsending.

Tracking: hvem har svart, med hvilke priser.
"""
from __future__ import annotations

import io
import json
import os
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None  # type: ignore

try:
    import requests
except ImportError:
    requests = None

try:
    from supabase import create_client, Client as SupabaseClient
except ImportError:
    create_client = None


PRICING_DIR = Path("qa_database/tender_pricing")
PRICING_DIR.mkdir(parents=True, exist_ok=True)
PACKAGES_LOG = PRICING_DIR / "pricing_packages.jsonl"


_SB = None


def _sb():
    global _SB
    if _SB is None and create_client is not None:
        url = os.environ.get("SUPABASE_URL", "").strip()
        key = (
            os.environ.get("SUPABASE_SERVICE_KEY", "").strip()
            or os.environ.get("SUPABASE_ANON_KEY", "").strip()
        )
        if url and key:
            try:
                _SB = create_client(url, key)
            except Exception:
                _SB = None
    return _SB


# ═════════════════════════════════════════════════════════════════
# PACKAGE GENERATION
# ═════════════════════════════════════════════════════════════════
def build_pricing_packages(
    pass3_data: Optional[Dict[str, Any]],
    pass1_results: List[Dict[str, Any]],
    config: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """
    Merge AI-recommended pricing packages with scope items per package
    from pass-1 extraction.
    """
    recommended = (pass3_data or {}).get("pricing_packages", []) if pass3_data else []
    project_packages = config.get("packages", [])

    # Build scope items index by package from pass-1
    scope_by_package: Dict[str, List[Dict[str, Any]]] = {}
    for r in pass1_results:
        ex = r.get("extraction") or {}
        for item in ex.get("scope_items", []) or []:
            pkg = (item.get("package") or "annet").lower()
            scope_by_package.setdefault(pkg, []).append({
                "description": item.get("description", ""),
                "quantity": item.get("quantity"),
                "unit": item.get("unit"),
                "source": r.get("filename"),
            })

    packages_out: List[Dict[str, Any]] = []
    seen = set()

    # First, use AI recommendations
    for rec in recommended:
        pkg = (rec.get("package") or "").lower()
        if not pkg or pkg in seen:
            continue
        seen.add(pkg)
        packages_out.append({
            "package": pkg,
            "send_to_external": rec.get("send_to_external", True),
            "rationale": rec.get("rationale", ""),
            "estimated_value_mnok": rec.get("estimated_value_mnok"),
            "key_specifications": rec.get("key_specifications", []),
            "open_questions": rec.get("open_questions", []),
            "suggested_suppliers_hint": rec.get("suggested_suppliers_hint", ""),
            "scope_items": scope_by_package.get(pkg, []),
            "source": "AI + parsing",
        })

    # Add any config-listed packages not in AI output
    for pkg in project_packages:
        pkg_low = pkg.lower()
        if pkg_low not in seen:
            seen.add(pkg_low)
            packages_out.append({
                "package": pkg_low,
                "send_to_external": True,
                "rationale": "Inkludert i konfigurasjonen — ikke eksplisitt vurdert av AI",
                "estimated_value_mnok": None,
                "key_specifications": [],
                "open_questions": [],
                "suggested_suppliers_hint": "",
                "scope_items": scope_by_package.get(pkg_low, []),
                "source": "Konfigurasjon",
            })

    return packages_out


# ═════════════════════════════════════════════════════════════════
# DOCX BUILDER — RFQ (Request For Quote) document per package
# ═════════════════════════════════════════════════════════════════
def build_rfq_docx(
    package: Dict[str, Any],
    project_info: Dict[str, Any],
    response_deadline_days: int = 10,
    contact_name: str = "",
    contact_email: str = "",
) -> Optional[bytes]:
    """Build a DOCX request-for-quote for a single package."""
    if Document is None:
        return None

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Header ──────────────────────────────────────────────────
    title = doc.add_paragraph()
    run = title.add_run("FORESPØRSEL OM PRIS")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x05, 0x6B, 0x8C)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    r = sub.add_run(f"Fagpakke: {package.get('package', '').title()}")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x38, 0x4B, 0x5E)

    doc.add_paragraph()

    # ── Project info box ────────────────────────────────────────
    proj = doc.add_paragraph()
    proj.add_run("Prosjekt: ").bold = True
    proj.add_run(project_info.get("p_name", "-") + "\n")
    proj.add_run("Adresse: ").bold = True
    proj.add_run(project_info.get("adresse", "-") + "\n")
    proj.add_run("Oppdragsgiver: ").bold = True
    proj.add_run(project_info.get("c_name") or project_info.get("company", "-"))

    doc.add_paragraph()

    # ── Deadline ────────────────────────────────────────────────
    deadline = datetime.now(timezone.utc) + timedelta(days=response_deadline_days)
    dl_para = doc.add_paragraph()
    dl_run = dl_para.add_run("Svarfrist: ")
    dl_run.bold = True
    dl_run.font.size = Pt(11)
    dl_para.add_run(deadline.strftime("%d.%m.%Y kl. 12:00"))

    doc.add_paragraph()

    # ── Rationale / scope intro ────────────────────────────────
    doc.add_heading("1. Bakgrunn for forespørselen", level=2)
    doc.add_paragraph(
        package.get("rationale")
        or f"Vi innhenter pristilbud for pakken {package.get('package', '').title()} "
        f"i forbindelse med ovennevnte prosjekt."
    )

    # ── Scope items ────────────────────────────────────────────
    doc.add_heading("2. Omfang / scope", level=2)
    scope_items = package.get("scope_items", [])
    if scope_items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Post"
        hdr[1].text = "Beskrivelse"
        hdr[2].text = "Mengde"
        hdr[3].text = "Enhet"
        for i, item in enumerate(scope_items[:50], 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = str(item.get("description", ""))[:300]
            row[2].text = str(item.get("quantity") or "-")
            row[3].text = str(item.get("unit") or "-")
    else:
        doc.add_paragraph(
            "Scope er beskrevet i vedlagt konkurransegrunnlag og beskrivelse. "
            "Tilbyder bes prise fullt omfang i henhold til vedlegg."
        )

    # ── Key specifications ─────────────────────────────────────
    specs = package.get("key_specifications", [])
    if specs:
        doc.add_heading("3. Sentrale krav og spesifikasjoner", level=2)
        for spec in specs[:20]:
            doc.add_paragraph(str(spec), style="List Bullet")

    # ── Open questions tilbyder must address ───────────────────
    questions = package.get("open_questions", [])
    if questions:
        doc.add_heading("4. Spørsmål tilbyder bes besvare", level=2)
        for q in questions[:15]:
            doc.add_paragraph(str(q), style="List Number")

    # ── Response format ────────────────────────────────────────
    doc.add_heading("5. Tilbudsformat", level=2)
    doc.add_paragraph("Tilbudet skal inneholde:")
    for line in [
        "Totalpris eks. mva., brutt ned på postnivå",
        "Enhetspriser for prising av evt. mengdeendringer",
        "Leveringstid og ressursdisponering",
        "Forutsetninger og eventuelle forbehold",
        "Referanser fra tilsvarende oppdrag (siste 3 år)",
        "Kontaktperson og signert tilbud",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Contact ────────────────────────────────────────────────
    doc.add_heading("6. Kontakt og innlevering", level=2)
    p = doc.add_paragraph()
    p.add_run("Forespørsel sendes til: ").bold = True
    p.add_run(f"{contact_name or '-'} — {contact_email or '-'}\n")
    p.add_run("Svarformat: ").bold = True
    p.add_run("PDF eller signert e-post\n")
    p.add_run("Forespørsels-ID: ").bold = True
    p.add_run(package.get("package_id", "-"))

    # ── Footer ─────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    frun = footer.add_run(
        "Dette underlaget er generert av Builtly Anbudskontroll. "
        "Mottaker er ikke forpliktet til å gi tilbud. Svar behandles konfidensielt."
    )
    frun.italic = True
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(0x78, 0x88, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════
# SEND VIA RESEND
# ═════════════════════════════════════════════════════════════════
def send_rfq_email(
    to_email: str,
    to_name: str,
    subject: str,
    body_markdown: str,
    docx_bytes: bytes,
    docx_filename: str,
    from_email: Optional[str] = None,
    from_name: Optional[str] = None,
    reply_to: Optional[str] = None,
) -> Dict[str, Any]:
    """Send RFQ via Resend API. Returns status dict."""
    api_key = os.environ.get("RESEND_API_KEY", "").strip()
    if not api_key or not requests:
        return {"status": "skipped", "reason": "RESEND_API_KEY mangler eller requests ikke installert"}

    import base64
    attachment_b64 = base64.b64encode(docx_bytes).decode("ascii")

    from_field = f"{from_name} <{from_email}>" if from_name and from_email else (
        from_email or os.environ.get("RESEND_FROM_EMAIL", "noreply@builtly.ai")
    )

    html_body = (
        "<div style='font-family: Inter, Arial, sans-serif; color: #1e2b3a; line-height: 1.55;'>"
        + body_markdown.replace("\n", "<br/>")
        + "</div>"
    )

    payload = {
        "from": from_field,
        "to": [to_email],
        "subject": subject,
        "html": html_body,
        "attachments": [
            {"filename": docx_filename, "content": attachment_b64}
        ],
    }
    if reply_to:
        payload["reply_to"] = reply_to

    try:
        resp = requests.post(
            "https://api.resend.com/emails",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            data=json.dumps(payload),
            timeout=20,
        )
        if resp.status_code in (200, 201, 202):
            return {"status": "sent", "resend_id": resp.json().get("id")}
        return {"status": "error", "http": resp.status_code, "body": resp.text[:500]}
    except Exception as e:
        return {"status": "error", "error": f"{type(e).__name__}: {e}"}


# ═════════════════════════════════════════════════════════════════
# TRACKING — persist package dispatch + responses
# ═════════════════════════════════════════════════════════════════
def log_package_dispatch(
    project_name: str,
    package: Dict[str, Any],
    recipient: Dict[str, str],
    send_result: Dict[str, Any],
    run_id: str,
) -> Dict[str, Any]:
    """Record a dispatched RFQ for tracking."""
    record = {
        "dispatch_id": str(uuid.uuid4()),
        "run_id": run_id,
        "project_name": project_name,
        "package_name": package.get("package"),
        "package_id": package.get("package_id"),
        "recipient_name": recipient.get("name"),
        "recipient_email": recipient.get("email"),
        "recipient_company": recipient.get("company"),
        "sent_at": datetime.now(timezone.utc).isoformat(),
        "status": send_result.get("status", "unknown"),
        "resend_id": send_result.get("resend_id"),
        "response_received": False,
        "response_price_mnok": None,
    }

    # Supabase
    sb = _sb()
    if sb is not None:
        try:
            sb.table("tender_pricing_dispatches").insert(record).execute()
        except Exception:
            pass

    # Local
    try:
        with open(PACKAGES_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False, default=str) + "\n")
    except Exception:
        pass

    return record


def load_dispatch_history(project_name: str) -> List[Dict[str, Any]]:
    """Return dispatch history for the project."""
    records: List[Dict[str, Any]] = []

    sb = _sb()
    if sb is not None:
        try:
            resp = (
                sb.table("tender_pricing_dispatches")
                .select("*")
                .eq("project_name", project_name)
                .order("sent_at", desc=True)
                .limit(100)
                .execute()
            )
            records.extend(resp.data or [])
        except Exception:
            pass

    if PACKAGES_LOG.exists():
        try:
            with open(PACKAGES_LOG, "r", encoding="utf-8") as f:
                for line in f.readlines()[-200:]:
                    try:
                        rec = json.loads(line)
                        if rec.get("project_name") == project_name:
                            if not any(r.get("dispatch_id") == rec.get("dispatch_id") for r in records):
                                records.append(rec)
                    except Exception:
                        continue
        except Exception:
            pass

    records.sort(key=lambda r: r.get("sent_at") or "", reverse=True)
    return records
